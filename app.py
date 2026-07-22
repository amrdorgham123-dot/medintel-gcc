"""
MedForsa GCC API — minimal real backend.
Run: uvicorn app:app --reload --port 8420
Docs auto-generated at: http://127.0.0.1:8420/docs
"""
from fastapi import FastAPI, HTTPException, Header, Depends, Request, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, Response, HTMLResponse
from pydantic import BaseModel, Field, EmailStr
from typing import Literal, Any
import sqlite3
import os
import json
import base64
import logging
import bcrypt
import jwt
import html as html_lib
import csv
import io
import secrets
import hashlib
import hmac
import smtplib
from email.mime.text import MIMEText
from datetime import datetime, timedelta

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
)
logger = logging.getLogger("medintel")

DB_PATH = os.path.join(os.path.dirname(__file__), "medintel.db")
FRONTEND_PATH = os.path.join(os.path.dirname(__file__), "frontend.html")
LANDING_PATH = os.path.join(os.path.dirname(__file__), "medintel-landing.html")
LANG_JS_PATH = os.path.join(os.path.dirname(__file__), "lang.js")

app = FastAPI(title="MedForsa GCC API", version="0.6")
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_methods=["*"], allow_headers=["*"])

def get_conn():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def log_action(conn, table, record_id, action, detail):
    conn.execute(
        "INSERT INTO audit_log (table_name, record_id, action, detail) VALUES (?,?,?,?)",
        (table, record_id, action, detail)
    )
    conn.commit()

class NewCompany(BaseModel):
    name: str = Field(..., min_length=1, max_length=200)
    headquarters: str = "Unknown"
    website: str = "Not provided"
    portfolio: str = "Not described yet"
    status_tag: Literal["covered", "unclear", "open"] = "unclear"

NO_CACHE_HEADERS = {
    "Cache-Control": "no-cache, no-store, must-revalidate",
    "Pragma": "no-cache",
    "Expires": "0",
}

@app.get("/")
def root():
    return FileResponse(LANDING_PATH, headers=NO_CACHE_HEADERS)

@app.get("/app")
def dashboard_page():
    return FileResponse(FRONTEND_PATH, headers=NO_CACHE_HEADERS)

@app.get("/lang.js")
def lang_js():
    return FileResponse(LANG_JS_PATH, media_type="application/javascript", headers=NO_CACHE_HEADERS)

@app.get("/snibe-ad.html")
def snibe_ad():
    return FileResponse(os.path.join(os.path.dirname(__file__), "snibe-ad.html"), media_type="text/html", headers=NO_CACHE_HEADERS)

@app.get("/snibe-biochem-ad.html")
def snibe_biochem_ad():
    return FileResponse(os.path.join(os.path.dirname(__file__), "snibe-biochem-ad.html"), media_type="text/html", headers=NO_CACHE_HEADERS)

@app.get("/doctor-ai")
def doctor_ai_page():
    return FileResponse(os.path.join(os.path.dirname(__file__), "doctor-ai.html"), media_type="text/html", headers=NO_CACHE_HEADERS)

@app.get("/admin")
def admin_page():
    return FileResponse(os.path.join(os.path.dirname(__file__), "admin.html"), media_type="text/html", headers=NO_CACHE_HEADERS)

@app.get("/lab-info")
def lab_info_page():
    return FileResponse(os.path.join(os.path.dirname(__file__), "lab-info.html"), media_type="text/html", headers=NO_CACHE_HEADERS)

# ---------------- SEO: server-rendered public pages, sitemap, robots.txt ----------------
# These return real, readable HTML with meta/OG/JSON-LD baked in server-side -- unlike
# /app (a JS-rendered SPA shell), a crawler or a link-preview bot sees actual content
# on first response, no JS execution required.

SITE_URL = os.environ.get("SITE_URL", "https://medintel-gcc.onrender.com")

def _seo_page_shell(title: str, description: str, canonical_path: str, body_html: str, json_ld: dict) -> str:
    title_e = html_lib.escape(title)
    desc_e = html_lib.escape(description)
    canonical = f"{SITE_URL}{canonical_path}"
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title_e} — MedForsa GCC</title>
<meta name="description" content="{desc_e}">
<link rel="canonical" href="{canonical}">
<meta property="og:title" content="{title_e}">
<meta property="og:description" content="{desc_e}">
<meta property="og:type" content="website">
<meta property="og:url" content="{canonical}">
<meta property="og:site_name" content="MedForsa GCC">
<meta name="twitter:card" content="summary">
<link rel="icon" href="/logo.svg" type="image/svg+xml">
<script type="application/ld+json">{json.dumps(json_ld)}</script>
<style>
  :root{{ --ink:#0A0E1C; --surface:#121A33; --paper:#EDEFF7; --muted:#8D95B3; --assay:#1FE0CE; --line:rgba(237,239,247,0.12); }}
  body{{ background:var(--ink); color:var(--paper); font-family:-apple-system,'Segoe UI',Tajawal,sans-serif; margin:0; line-height:1.65; }}
  header{{ padding:20px 24px; border-bottom:1px solid var(--line); }}
  header a{{ color:var(--assay); text-decoration:none; font-weight:600; font-size:14px; }}
  main{{ max-width:760px; margin:0 auto; padding:36px 24px 80px; }}
  h1{{ font-size:26px; margin:0 0 6px; }}
  .sub{{ color:var(--muted); font-size:14px; margin-bottom:28px; }}
  .card{{ background:var(--surface); border:1px solid var(--line); border-radius:14px; padding:20px 22px; margin-bottom:16px; }}
  .card h2{{ font-size:15px; margin:0 0 10px; color:var(--assay); }}
  .card p{{ margin:0 0 8px; font-size:14.5px; }}
  .row{{ display:flex; justify-content:space-between; gap:12px; padding:6px 0; border-bottom:1px solid var(--line); font-size:13.5px; }}
  .row:last-child{{ border-bottom:none; }}
  .row .label{{ color:var(--muted); }}
  .cta{{ display:inline-block; margin-top:8px; background:var(--assay); color:var(--ink); padding:10px 18px; border-radius:999px; text-decoration:none; font-weight:600; font-size:14px; }}
  a{{ color:var(--assay); }}
</style>
</head>
<body>
<header><a href="/">← MedForsa GCC</a></header>
<main>{body_html}</main>
</body>
</html>"""

@app.get("/company/{slug}", response_class=HTMLResponse)
def company_seo_page(slug: str):
    conn = get_conn()
    m = conn.execute("SELECT * FROM manufacturers WHERE slug = ? AND is_published = 1", (slug,)).fetchone()
    if not m:
        conn.close()
        raise HTTPException(status_code=404, detail="Company not found")
    m = dict(m)
    products = conn.execute("SELECT product_name, product_type FROM products WHERE manufacturer_id = ? LIMIT 20", (m["id"],)).fetchall()
    distributors = conn.execute("""
        SELECT d.name, d.country FROM company_distributors cd
        JOIN distributors d ON d.id = cd.distributor_id WHERE cd.manufacturer_id = ?
    """, (m["id"],)).fetchall()
    conn.close()

    description = f"{m['name']} -- {m['category']} manufacturer. {m.get('portfolio') or ''}"[:250]
    body = f"""
<h1>{html_lib.escape(m['name'])}</h1>
<p class="sub">{html_lib.escape(m['category'] or '')} · {html_lib.escape(m.get('origin') or '')} · Headquarters: {html_lib.escape(m.get('headquarters') or 'Not listed')}</p>
<div class="card">
  <h2>Overview</h2>
  <p>{html_lib.escape(m.get('portfolio') or 'Portfolio details not yet published.')}</p>
  <div class="row"><span class="label">Saudi Arabia distribution status</span><span>{html_lib.escape(m.get('ksa_status') or 'Unclear')}</span></div>
  <div class="row"><span class="label">Website</span><span>{html_lib.escape(m.get('website') or 'Not listed')}</span></div>
</div>
{"<div class='card'><h2>Products (" + str(len(products)) + ")</h2>" + "".join(f"<div class='row'><span>{html_lib.escape(p['product_name'])}</span><span class='label'>{html_lib.escape(p['product_type'] or '')}</span></div>" for p in products) + "</div>" if products else ""}
{"<div class='card'><h2>Confirmed KSA distributors</h2>" + "".join(f"<div class='row'><span>{html_lib.escape(d['name'])}</span><span class='label'>{html_lib.escape(d['country'] or '')}</span></div>" for d in distributors) + "</div>" if distributors else ""}
<a class="cta" href="/app">Explore the full platform →</a>
"""
    json_ld = {
        "@context": "https://schema.org", "@type": "Organization",
        "name": m["name"], "description": description,
        "url": f"{SITE_URL}/company/{slug}",
    }
    return _seo_page_shell(m["name"], description, f"/company/{slug}", body, json_ld)

@app.get("/lab-test/{slug}", response_class=HTMLResponse)
def lab_test_seo_page(slug: str):
    conn = get_conn()
    t = conn.execute("SELECT * FROM lab_tests WHERE slug = ? AND is_published = 1", (slug,)).fetchone()
    conn.close()
    if not t:
        raise HTTPException(status_code=404, detail="Lab test not found")
    t = dict(t)
    ranges = json.loads(t.get("reference_ranges_json") or "[]")
    description = (t.get("purpose_en") or f"{t['name_en']} -- reference ranges, clinical significance, and interpretation.")[:250]
    body = f"""
<h1>{html_lib.escape(t['name_en'])}</h1>
<p class="sub">{html_lib.escape(t['category'] or '')}</p>
<div class="card">
  <h2>Purpose</h2>
  <p>{html_lib.escape(t.get('purpose_en') or '')}</p>
</div>
{"<div class='card'><h2>Reference ranges</h2>" + "".join(f"<div class='row'><span>{html_lib.escape(r.get('population',''))} {html_lib.escape(r.get('parameter',''))}</span><span class='label'>{html_lib.escape(str(r.get('range','')))}</span></div>" for r in ranges) + "</div>" if ranges else ""}
{"<div class='card'><h2>Clinical significance</h2><p>" + html_lib.escape(t.get('clinical_significance_en') or '') + "</p></div>" if t.get('clinical_significance_en') else ""}
<a class="cta" href="/lab-info">Browse the full Lab Info library →</a>
"""
    json_ld = {
        "@context": "https://schema.org", "@type": "MedicalTest",
        "name": t["name_en"], "description": description,
        "url": f"{SITE_URL}/lab-test/{slug}",
    }
    return _seo_page_shell(t["name_en"], description, f"/lab-test/{slug}", body, json_ld)

@app.get("/sitemap.xml")
def sitemap_xml():
    conn = get_conn()
    companies = conn.execute("SELECT slug, created_at FROM manufacturers WHERE is_published = 1 AND slug IS NOT NULL").fetchall()
    tests = conn.execute("SELECT slug, updated_at FROM lab_tests WHERE is_published = 1").fetchall()
    conn.close()
    urls = [f"<url><loc>{SITE_URL}/</loc><changefreq>daily</changefreq></url>",
            f"<url><loc>{SITE_URL}/lab-info</loc><changefreq>weekly</changefreq></url>"]
    for c in companies:
        urls.append(f"<url><loc>{SITE_URL}/company/{c['slug']}</loc><changefreq>monthly</changefreq></url>")
    for t in tests:
        urls.append(f"<url><loc>{SITE_URL}/lab-test/{t['slug']}</loc><changefreq>monthly</changefreq></url>")
    xml = '<?xml version="1.0" encoding="UTF-8"?>\n<urlset xmlns="http://www.sitemaps.org/schemas/sitemap/0.9">\n' + "\n".join(urls) + "\n</urlset>"
    return Response(content=xml, media_type="application/xml")

@app.get("/robots.txt")
def robots_txt():
    body = f"User-agent: *\nAllow: /\nDisallow: /admin\nDisallow: /app\nSitemap: {SITE_URL}/sitemap.xml\n"
    return Response(content=body, media_type="text/plain")

# ---------------- Phase F: growth infrastructure (demo requests, public pricing page) ----------------

class DemoRequest(BaseModel):
    full_name: str = Field(..., min_length=1, max_length=150)
    company_name: str = Field(..., min_length=1, max_length=150)
    email: EmailStr
    phone: str | None = None
    message: str | None = None
    plan_interest: Literal["trial", "pro", "enterprise"] | None = None

@app.post("/demo-requests")
def submit_demo_request(payload: DemoRequest):
    """Public lead-capture endpoint -- this is what actually generates business
    before Stripe/Moyasar credentials exist. No login required: this is the
    landing page's 'Request a Demo' form."""
    conn = get_conn()
    conn.execute(
        "INSERT INTO demo_requests (full_name, company_name, email, phone, message, plan_interest) VALUES (?,?,?,?,?,?)",
        (payload.full_name, payload.company_name, payload.email, payload.phone, payload.message, payload.plan_interest)
    )
    conn.commit()
    conn.close()
    logger.info(f"New demo request: {payload.full_name} ({payload.company_name}) -- {payload.email}")
    return {"status": "received", "message": "Thanks -- we'll be in touch shortly."}

@app.get("/pricing", response_class=HTMLResponse)
def pricing_page():
    """Public, server-rendered pricing page (SEO-indexable, matches Phase D's
    page shell) -- shows the real plan data from /subscription/plans rather
    than a separately-maintained copy."""
    plans_html = ""
    for key, plan in SUBSCRIPTION_PLANS.items():
        price = f"{plan['price_sar']} SAR" if plan["price_sar"] else "Contact us"
        features = "".join(f"<div class='row'><span>{html_lib.escape(f)}</span></div>" for f in plan["features"])
        plans_html += f"""
        <div class="card">
          <h2>{html_lib.escape(plan['name'])}</h2>
          <p style="font-size:22px;font-weight:600;color:var(--paper);margin:0 0 4px">{price}</p>
          <p style="color:var(--muted);font-size:12.5px;margin:0 0 14px">{html_lib.escape(plan['billing'])}</p>
          {features}
        </div>"""
    body = f"""
<h1>Plans and pricing</h1>
<p class="sub">Full platform access, transparent pricing. Online self-serve checkout is coming soon -- for now, request a demo below and we'll set up your account directly.</p>
{plans_html}
<div class="card" id="demo-form">
  <h2>Request a demo</h2>
  <form onsubmit="event.preventDefault(); submitDemoRequest(this);" style="display:flex;flex-direction:column;gap:10px;">
    <input name="full_name" required placeholder="Full name" style="padding:10px;border-radius:8px;border:1px solid var(--line);background:var(--surface);color:var(--paper)">
    <input name="company_name" required placeholder="Company name" style="padding:10px;border-radius:8px;border:1px solid var(--line);background:var(--surface);color:var(--paper)">
    <input name="email" type="email" required placeholder="Email" style="padding:10px;border-radius:8px;border:1px solid var(--line);background:var(--surface);color:var(--paper)">
    <input name="phone" placeholder="Phone (optional)" style="padding:10px;border-radius:8px;border:1px solid var(--line);background:var(--surface);color:var(--paper)">
    <button type="submit" class="cta" style="border:none;cursor:pointer">Request a demo</button>
    <p id="demo-form-msg" style="font-size:13px;color:var(--assay);"></p>
  </form>
</div>
<script>
async function submitDemoRequest(form){{
  const data = Object.fromEntries(new FormData(form));
  const msgEl = document.getElementById('demo-form-msg');
  try {{
    const res = await fetch('/demo-requests', {{method:'POST', headers:{{'Content-Type':'application/json'}}, body: JSON.stringify(data)}});
    if(res.ok){{ msgEl.textContent = "Thanks -- we'll be in touch shortly."; form.reset(); }}
    else {{ msgEl.textContent = 'Something went wrong -- please try again.'; }}
  }} catch(e){{ msgEl.textContent = 'Could not reach the server.'; }}
}}
</script>
"""
    json_ld = {"@context": "https://schema.org", "@type": "PriceSpecification", "name": "MedForsa GCC Plans"}
    return _seo_page_shell("Plans and pricing", "MedForsa GCC subscription plans -- Trial, Pro, and Enterprise access to IVD and blood bank market intelligence across Saudi Arabia and the GCC.", "/pricing", body, json_ld)

@app.get("/lab-tests")
def lab_tests_page_legacy():
    # Legacy URL kept working -- the page and feature are now branded "Lab Info".
    return FileResponse(os.path.join(os.path.dirname(__file__), "lab-info.html"), media_type="text/html", headers=NO_CACHE_HEADERS)

@app.get("/compatibility-wizard")
def compatibility_wizard_page():
    return FileResponse(os.path.join(os.path.dirname(__file__), "compatibility-wizard.html"), media_type="text/html", headers=NO_CACHE_HEADERS)

@app.get("/tco-calculator")
def tco_calculator_page():
    return FileResponse(os.path.join(os.path.dirname(__file__), "tco-calculator.html"), media_type="text/html", headers=NO_CACHE_HEADERS)

@app.get("/logo.svg")
def logo():
    return FileResponse(os.path.join(os.path.dirname(__file__), "logo.svg"), media_type="image/svg+xml")

@app.get("/api/status")
def status():
    return {"service": "MedForsa GCC API", "status": "running", "note": "This is a real local backend, not a hosted service."}

class ChatMessage(BaseModel):
    role: Literal["user", "assistant"]
    content: Any

class ChatRequest(BaseModel):
    messages: list[ChatMessage]
    system: str = ""

def call_anthropic(system: str, messages: list):
    import json
    import urllib.request
    import urllib.error

    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="ANTHROPIC_API_KEY is not set on the server. Add it in Render > Environment.")

    # Strip accidental whitespace/quotes that sometimes sneak in when pasting
    # the key into Render's Environment UI — a common source of "invalid x-api-key".
    api_key = api_key.strip().strip('"').strip("'")

    # Safe diagnostic (never logs the full key): confirms what actually reached
    # the server without exposing the secret in logs.
    logger.info(
        f"Anthropic key loaded: length={len(api_key)}, "
        f"prefix={api_key[:7]}, suffix={api_key[-4:] if len(api_key) >= 4 else '****'}"
    )
    if not api_key.startswith("sk-ant-"):
        logger.warning("ANTHROPIC_API_KEY does not start with 'sk-ant-' — this is likely the wrong value.")

    body = json.dumps({
        "model": "claude-sonnet-5",
        "max_tokens": 1500,
        "system": system,
        "messages": messages,
    }).encode("utf-8")

    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=body,
        headers={
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=45) as resp:
            return json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        detail = e.read().decode("utf-8")
        logger.error(f"Anthropic API error {e.code}: {detail}")
        raise HTTPException(status_code=e.code, detail=detail)
    except Exception as e:
        logger.error(f"Chat proxy error: {e}")
        raise HTTPException(status_code=502, detail=str(e))

@app.post("/api/chat")
def chat_proxy(payload: ChatRequest):
    return call_anthropic(payload.system, [m.dict() for m in payload.messages])

@app.get("/companies")
def list_companies(status: str | None = None, q: str | None = None, category: str | None = None,
                    origin: str | None = None, distributor_id: int | None = None,
                    limit: int = 50, offset: int = 0, include_pending: bool = False):
    if limit < 1 or limit > 200:
        raise HTTPException(status_code=400, detail="limit must be between 1 and 200")
    if offset < 0:
        raise HTTPException(status_code=400, detail="offset must be >= 0")
    conn = get_conn()
    query = "SELECT * FROM manufacturers WHERE 1=1"
    params = []
    if not include_pending:
        query += " AND is_published = 1"
    if status:
        query += " AND status_tag = ?"
        params.append(status)
    if category:
        query += " AND category = ?"
        params.append(category)
    if origin:
        query += " AND origin = ?"
        params.append(origin)
    if distributor_id:
        query += " AND id IN (SELECT manufacturer_id FROM company_distributors WHERE distributor_id = ?)"
        params.append(distributor_id)
    if q:
        query += " AND (name LIKE ? OR portfolio LIKE ? OR ksa_status LIKE ?)"
        like = f"%{q}%"
        params += [like, like, like]
    total = conn.execute(f"SELECT count(*) c FROM ({query})", params).fetchone()["c"]
    query += " ORDER BY id LIMIT ? OFFSET ?"
    params += [limit, offset]
    rows = conn.execute(query, params).fetchall()
    conn.close()
    logger.info(f"GET /companies status={status} category={category} origin={origin} q={q} -> {len(rows)}/{total} results")
    return {
        "results": [dict(r) for r in rows],
        "total": total,
        "limit": limit,
        "offset": offset,
        "has_more": offset + len(rows) < total,
    }

@app.get("/review-queue")
def review_queue():
    """Companies added but not yet approved by a human reviewer -- the real, bounded
    version of 'nothing gets published without review'."""
    conn = get_conn()
    rows = conn.execute("SELECT * FROM manufacturers WHERE is_published = 0 ORDER BY id").fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/search")
def universal_search(q: str, limit_per_type: int = 5):
    """Single search box across every real content type on the platform:
    manufacturers, products, technologies, distributors, and lab tests.
    Ranked within each type (exact match > starts-with > contains), and
    returns a flat ranked list the frontend can render grouped or ungrouped.
    Public endpoint -- doesn't require login, matching the rest of the
    directory/lab-info browsing experience."""
    q = (q or "").strip()
    if len(q) < 2:
        return {"query": q, "results": []}
    like = f"%{q}%"
    conn = get_conn()
    results = []

    def rank(name: str) -> int:
        n = name.lower()
        ql = q.lower()
        if n == ql:
            return 0
        if n.startswith(ql):
            return 1
        return 2

    mfrs = conn.execute(
        "SELECT id, name, category FROM manufacturers WHERE is_published = 1 AND name LIKE ? LIMIT ?",
        (like, limit_per_type)
    ).fetchall()
    for m in mfrs:
        results.append({"type": "company", "id": m["id"], "title": m["name"], "subtitle": m["category"],
                         "url": f"/app#company-{m['id']}", "rank": rank(m["name"])})

    products = conn.execute(
        """SELECT p.id, p.product_name, m.name as mfr_name FROM products p
           LEFT JOIN manufacturers m ON p.manufacturer_id = m.id
           WHERE p.product_name LIKE ? LIMIT ?""",
        (like, limit_per_type)
    ).fetchall()
    for p in products:
        results.append({"type": "product", "id": p["id"], "title": p["product_name"], "subtitle": p["mfr_name"],
                         "url": f"/app#product-{p['id']}", "rank": rank(p["product_name"])})

    techs = conn.execute(
        "SELECT id, name, description FROM technologies WHERE name LIKE ? LIMIT ?",
        (like, limit_per_type)
    ).fetchall()
    for t in techs:
        subtitle = (t["description"] or "")[:80]
        results.append({"type": "technology", "id": t["id"], "title": t["name"], "subtitle": subtitle,
                         "url": f"/app#technology-{t['id']}", "rank": rank(t["name"])})

    dists = conn.execute(
        "SELECT id, name, country, represents FROM distributors WHERE name LIKE ? OR represents LIKE ? LIMIT ?",
        (like, like, limit_per_type)
    ).fetchall()
    for d in dists:
        results.append({"type": "distributor", "id": d["id"], "title": d["name"], "subtitle": d["country"],
                         "url": f"/app#distributor-{d['id']}", "rank": rank(d["name"])})

    tests = conn.execute(
        "SELECT id, slug, name_en, category FROM lab_tests WHERE is_published = 1 AND (name_en LIKE ? OR aliases LIKE ?) LIMIT ?",
        (like, like, limit_per_type)
    ).fetchall()
    for t in tests:
        results.append({"type": "lab_test", "id": t["id"], "title": t["name_en"], "subtitle": t["category"],
                         "url": f"/lab-tests-detail/{t['slug']}", "rank": rank(t["name_en"])})

    conn.close()
    results.sort(key=lambda r: (r["rank"], r["title"]))
    return {"query": q, "results": results}

@app.post("/companies/{company_id}/approve")
def approve_company(company_id: int):
    conn = get_conn()
    existing = conn.execute("SELECT name, is_published FROM manufacturers WHERE id = ?", (company_id,)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Company not found")
    conn.execute("UPDATE manufacturers SET is_published = 1 WHERE id = ?", (company_id,))
    conn.commit()
    log_action(conn, "manufacturers", company_id, "approve", f"Approved and published: {existing['name']}")
    logger.info(f"POST /companies/{company_id}/approve -> published {existing['name']}")
    conn.close()
    return {"id": company_id, "status": "approved_and_published"}

@app.delete("/companies/{company_id}/reject")
def reject_company(company_id: int):
    conn = get_conn()
    existing = conn.execute("SELECT name FROM manufacturers WHERE id = ? AND is_published = 0", (company_id,)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Pending company not found (already published or never existed -- use DELETE /companies/{id} for published records)")
    conn.execute("DELETE FROM manufacturers WHERE id = ?", (company_id,))
    conn.commit()
    log_action(conn, "manufacturers", company_id, "reject", f"Rejected during review: {existing['name']}")
    logger.info(f"DELETE /companies/{company_id}/reject -> rejected {existing['name']}")
    conn.close()
    return {"id": company_id, "status": "rejected_and_removed"}

@app.get("/categories")
def list_categories():
    conn = get_conn()
    rows = conn.execute("SELECT category, count(*) as c FROM manufacturers GROUP BY category").fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/origins")
def list_origins():
    conn = get_conn()
    rows = conn.execute("SELECT origin, count(*) as c FROM manufacturers GROUP BY origin").fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/attieh-portfolio")
def attieh_portfolio():
    """Manufacturers Attieh Medico is directly and confirmedly awarded to represent, per official NUPCO tender records."""
    conn = get_conn()
    rows = conn.execute("SELECT * FROM attieh_portfolio ORDER BY manufacturer").fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/tenders")
def list_tenders():
    """Real tender references from the two NUPCO documents provided -- not a live tender feed."""
    conn = get_conn()
    rows = conn.execute("SELECT * FROM tenders ORDER BY tender_ref").fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/hospitals")
def list_hospitals():
    """Real installed-equipment data for major KSA government hospital networks, sourced from NUPCO tender NPT0058/25."""
    conn = get_conn()
    rows = conn.execute("SELECT * FROM hospitals ORDER BY abbreviation").fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/products")
def list_products(manufacturer_id: int | None = None, origin: str | None = None,
                   department: str | None = None, technology: str | None = None,
                   distributor_id: int | None = None):
    conn = get_conn()
    query = """SELECT p.*, m.name as manufacturer_name, m.category, m.confidence_tier, m.origin
               FROM products p JOIN manufacturers m ON m.id = p.manufacturer_id WHERE 1=1"""
    params = []
    if manufacturer_id:
        query += " AND p.manufacturer_id = ?"
        params.append(manufacturer_id)
    if origin:
        query += " AND m.origin = ?"
        params.append(origin)
    if department:
        query += " AND p.department = ?"
        params.append(department)
    if distributor_id:
        query += " AND m.id IN (SELECT manufacturer_id FROM company_distributors WHERE distributor_id = ?)"
        params.append(distributor_id)
    if technology:
        query += """ AND m.id IN (
            SELECT mt.manufacturer_id FROM manufacturer_technology mt
            JOIN technologies t ON t.id = mt.technology_id WHERE t.name LIKE ?)"""
        params.append(f"%{technology}%")
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/distributor-products/{distributor_name}")
def distributor_products(distributor_name: str):
    conn = get_conn()
    # Use LIKE matching since the same real-world distributor sometimes appears under
    # slightly different name formats across different source documents (e.g. "Samir
    # Group" vs "Samir Trading & Marketing") -- exact match would silently miss real links.
    key = distributor_name.split(" (")[0].split(",")[0].strip()
    rows = conn.execute(
        """SELECT dp.*, p.product_name, p.product_type, p.description, m.name as manufacturer_name
           FROM distributor_products dp
           JOIN products p ON p.id = dp.product_id
           JOIN manufacturers m ON m.id = p.manufacturer_id
           WHERE dp.distributor_name LIKE ? OR ? LIKE '%' || dp.distributor_name || '%'""",
        (f"%{key}%", distributor_name)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/opportunities-by-category")
def opportunities_by_category():
    """Honest per-category opportunity coverage -- reports zero, not a fabricated lead, when every
    researched company in a category already has confirmed KSA coverage."""
    conn = get_conn()
    categories = [r["category"] for r in conn.execute("SELECT DISTINCT category FROM manufacturers").fetchall()]
    result = []
    for cat in categories:
        total = conn.execute("SELECT count(*) c FROM manufacturers WHERE category = ?", (cat,)).fetchone()["c"]
        open_or_unclear = conn.execute(
            "SELECT count(*) c FROM manufacturers WHERE category = ? AND status_tag IN ('open','unclear')", (cat,)
        ).fetchone()["c"]
        result.append({
            "category": cat,
            "companies_researched": total,
            "companies_with_open_question": open_or_unclear,
            "note": "All researched companies in this category already have confirmed KSA coverage." if open_or_unclear == 0
                    else f"{open_or_unclear} of {total} researched companies have an unresolved KSA distribution question."
        })
    conn.close()
    return result

@app.get("/companies/{company_id}/intelligence-score")
def intelligence_score(company_id: int):
    """A transparent, explainable score computed from real fields already in the database.
    Not a black box -- every point is shown with its source."""
    conn = get_conn()
    company = conn.execute("SELECT * FROM manufacturers WHERE id = ?", (company_id,)).fetchone()
    if not company:
        conn.close()
        raise HTTPException(status_code=404, detail="Company not found")
    n_products = conn.execute("SELECT count(*) c FROM products WHERE manufacturer_id = ?", (company_id,)).fetchone()["c"]
    n_technologies = conn.execute(
        "SELECT count(*) c FROM manufacturer_technology WHERE manufacturer_id = ?", (company_id,)
    ).fetchone()["c"]
    n_evidence = conn.execute("SELECT count(*) c FROM evidence WHERE manufacturer_id = ?", (company_id,)).fetchone()["c"]
    n_events = conn.execute("SELECT count(*) c FROM market_events WHERE manufacturer_id = ?", (company_id,)).fetchone()["c"]
    conn.close()

    tier_points = {"gold": 30, "silver": 18, "bronze": 6}
    confidence_points = tier_points.get(company["confidence_tier"], 0)
    product_points = min(n_products * 5, 20)          # cap at 20 (4+ products)
    technology_points = min(n_technologies * 5, 15)    # cap at 15 (3+ technologies)
    evidence_points = min(n_evidence * 5, 20)          # cap at 20 (4+ evidence entries)
    event_points = min(n_events * 5, 10)               # cap at 10 (2+ events)
    distribution_points = 5 if company["status_tag"] == "covered" else 0  # a resolved status, not "open is better"

    total = confidence_points + product_points + technology_points + evidence_points + event_points + distribution_points

    return {
        "company": company["name"],
        "score": total,
        "max_possible": 100,
        "breakdown": {
            "confidence_tier": {"points": confidence_points, "max": 30, "basis": f"tier = {company['confidence_tier']}"},
            "product_portfolio": {"points": product_points, "max": 20, "basis": f"{n_products} products on file"},
            "technology_diversity": {"points": technology_points, "max": 15, "basis": f"{n_technologies} technology links"},
            "evidence_coverage": {"points": evidence_points, "max": 20, "basis": f"{n_evidence} evidence entries"},
            "market_event_history": {"points": event_points, "max": 10, "basis": f"{n_events} logged events"},
            "ksa_status_resolved": {"points": distribution_points, "max": 5, "basis": company["status_tag"]},
        },
        "note": "Every point above is computed from real database counts -- not an estimate or an AI judgment call."
    }

@app.get("/companies/{company_id}/competitors")
def company_competitors(company_id: int):
    """Competitors derived from shared technology links -- not researched separately, purely
    computed from data already in the database."""
    conn = get_conn()
    company = conn.execute("SELECT * FROM manufacturers WHERE id = ?", (company_id,)).fetchone()
    if not company:
        conn.close()
        raise HTTPException(status_code=404, detail="Company not found")
    rows = conn.execute(
        """SELECT DISTINCT m2.id, m2.name, m2.category, t.name as shared_technology
           FROM manufacturer_technology mt1
           JOIN manufacturer_technology mt2 ON mt1.technology_id = mt2.technology_id AND mt2.manufacturer_id != mt1.manufacturer_id
           JOIN manufacturers m2 ON m2.id = mt2.manufacturer_id
           JOIN technologies t ON t.id = mt1.technology_id
           WHERE mt1.manufacturer_id = ?""",
        (company_id,)
    ).fetchall()
    conn.close()
    if not rows:
        return {"company": company["name"], "competitors": [], "note": "No shared-technology links found -- either this company has no technology links recorded, or no other company shares one."}
    return {"company": company["name"], "competitors": [dict(r) for r in rows],
            "note": "Derived purely from shared Technology links already in the database -- not independently researched competitor analysis."}

@app.get("/market-events")
def market_events():
    conn = get_conn()
    rows = conn.execute(
        """SELECT e.*, m.name as manufacturer_name FROM market_events e
           JOIN manufacturers m ON m.id = e.manufacturer_id ORDER BY e.event_date DESC"""
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/companies/{company_id}")
def get_company(company_id: int):
    conn = get_conn()
    row = conn.execute("SELECT * FROM manufacturers WHERE id = ?", (company_id,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Company not found")
    company = dict(row)
    techs = conn.execute(
        """SELECT t.name, t.description FROM technologies t
           JOIN manufacturer_technology mt ON mt.technology_id = t.id
           WHERE mt.manufacturer_id = ?""", (company_id,)
    ).fetchall()
    company["technologies"] = [dict(t) for t in techs]
    dists = conn.execute(
        """SELECT d.* FROM distributors d
           JOIN company_distributors cd ON cd.distributor_id = d.id
           WHERE cd.manufacturer_id = ?""", (company_id,)
    ).fetchall()
    company["distributors"] = [dict(d) for d in dists]
    opp = conn.execute("SELECT * FROM opportunities WHERE manufacturer_id = ?", (company_id,)).fetchone()
    company["opportunity"] = dict(opp) if opp else None
    evidence = conn.execute("SELECT * FROM evidence WHERE manufacturer_id = ?", (company_id,)).fetchall()
    company["evidence"] = [dict(e) for e in evidence]
    products = conn.execute("SELECT * FROM products WHERE manufacturer_id = ?", (company_id,)).fetchall()
    company["products"] = [dict(p) for p in products]
    conn.close()
    return company

@app.get("/companies/{company_id}/evidence")
def company_evidence(company_id: int):
    conn = get_conn()
    rows = conn.execute("SELECT * FROM evidence WHERE manufacturer_id = ?", (company_id,)).fetchall()
    conn.close()
    if not rows:
        return {"company_id": company_id, "evidence": [], "note": "No structured evidence recorded yet for this company."}
    return {"company_id": company_id, "evidence": [dict(r) for r in rows]}

@app.get("/data-quality")
def data_quality():
    conn = get_conn()
    total = conn.execute("SELECT count(*) c FROM manufacturers").fetchone()["c"]
    tier_map = {"gold": 100, "silver": 65, "bronze": 30}
    rows = conn.execute("SELECT confidence_tier FROM manufacturers").fetchall()
    tier_counts = {"gold": 0, "silver": 0, "bronze": 0}
    total_score = 0
    for r in rows:
        t = r["confidence_tier"]
        tier_counts[t] = tier_counts.get(t, 0) + 1
        total_score += tier_map.get(t, 0)
    incomplete = conn.execute(
        "SELECT count(*) c FROM manufacturers WHERE sources LIKE '%None yet%' OR sources = '' OR sources IS NULL"
    ).fetchone()["c"]
    needs_review = conn.execute(
        "SELECT count(*) c FROM manufacturers WHERE status_tag = 'unclear' OR confidence_tier = 'bronze'"
    ).fetchone()["c"]
    with_evidence = conn.execute(
        "SELECT count(DISTINCT manufacturer_id) c FROM evidence"
    ).fetchone()["c"]
    latest = conn.execute("SELECT MAX(created_at) c FROM manufacturers").fetchone()["c"]
    conn.close()
    return {
        "total_companies": total,
        "average_confidence_score": round(total_score / total, 1) if total else 0,
        "tier_breakdown": tier_counts,
        "companies_with_no_sources_recorded": incomplete,
        "companies_needing_review": needs_review,
        "companies_with_structured_evidence": with_evidence,
        "companies_missing_structured_evidence": total - with_evidence,
        "last_record_added": latest,
        "note": "Every number above is computed live from the actual database, not estimated.",
    }

@app.post("/companies")
def add_company(company: NewCompany):
    conn = get_conn()
    cur = conn.execute(
        """INSERT INTO manufacturers (name, headquarters, website, portfolio, ksa_status, status_tag, opportunity_note, confidence_tier, sources, is_published)
           VALUES (?,?,?,?,?,?,?,?,?,0)""",
        (company.name, company.headquarters, company.website, company.portfolio,
         "Manually entered — not yet verified against an official source", company.status_tag,
         "Not yet scored", "bronze", "None yet — needs verification")
    )
    conn.commit()
    new_id = cur.lastrowid
    log_action(conn, "manufacturers", new_id, "insert", f"Submitted for review: {company.name}")
    logger.info(f"POST /companies -> submitted id={new_id} name={company.name} (pending review)")
    conn.close()
    return {"id": new_id, "status": "pending_review", "note": "Not published yet -- goes to the review queue. Call POST /companies/{id}/approve to publish it, or DELETE /companies/{id}/reject to discard it."}

class UpdateCompany(BaseModel):
    headquarters: str | None = None
    website: str | None = None
    portfolio: str | None = None
    status_tag: Literal["covered", "unclear", "open"] | None = None

@app.put("/companies/{company_id}")
def update_company(company_id: int, updates: UpdateCompany):
    conn = get_conn()
    existing = conn.execute("SELECT * FROM manufacturers WHERE id = ?", (company_id,)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Company not found")
    fields, params = [], []
    for field, value in updates.model_dump(exclude_unset=True).items():
        fields.append(f"{field} = ?")
        params.append(value)
    if not fields:
        conn.close()
        raise HTTPException(status_code=400, detail="No fields provided to update")
    params.append(company_id)
    conn.execute(f"UPDATE manufacturers SET {', '.join(fields)} WHERE id = ?", params)
    conn.commit()
    log_action(conn, "manufacturers", company_id, "update", f"Fields changed: {', '.join(updates.model_dump(exclude_unset=True).keys())}")
    logger.info(f"PUT /companies/{company_id} -> updated fields: {list(updates.model_dump(exclude_unset=True).keys())}")
    conn.close()
    return {"id": company_id, "status": "updated"}

@app.delete("/companies/{company_id}")
def delete_company(company_id: int):
    conn = get_conn()
    existing = conn.execute("SELECT name FROM manufacturers WHERE id = ?", (company_id,)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Company not found")
    conn.execute("DELETE FROM manufacturers WHERE id = ?", (company_id,))
    conn.commit()
    log_action(conn, "manufacturers", company_id, "delete", f"Deleted: {existing['name']}")
    logger.warning(f"DELETE /companies/{company_id} -> removed {existing['name']}")
    conn.close()
    return {"id": company_id, "status": "deleted"}

JWT_SECRET = os.environ.get("JWT_SECRET")
if not JWT_SECRET:
    JWT_SECRET = "insecure-local-dev-secret-change-me"
    logger.warning("JWT_SECRET env var not set -- using an insecure local-dev default. Set a real secret before deploying anywhere reachable by others.")
JWT_ALGORITHM = "HS256"
JWT_EXPIRE_HOURS = 24 * 365  # ~1 year -- effectively "stay logged in until you log out"

def create_access_token(user_id: int, email: str) -> str:
    payload = {"sub": str(user_id), "email": email, "exp": datetime.utcnow() + timedelta(hours=JWT_EXPIRE_HOURS)}
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)

def get_current_user(authorization: str | None = Header(default=None)) -> dict:
    """Real per-account auth. Every tenant-scoped endpoint (leads, watchlist,
    opportunities, daily-briefing) depends on this instead of a shared password
    -- each user only ever sees their own data."""
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid Authorization header")
    token = authorization.split(" ", 1)[1]
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Session expired, please log in again")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid session token")
    conn = get_conn()
    user = conn.execute("SELECT * FROM users WHERE id = ? AND is_active = 1", (payload["sub"],)).fetchone()
    conn.close()
    if not user:
        raise HTTPException(status_code=401, detail="Account not found or deactivated")
    return dict(user)

def public_user_fields(user: dict) -> dict:
    return {k: user[k] for k in ("id", "email", "company_name", "full_name", "role")}

def require_admin(current_user: dict = Depends(get_current_user)) -> dict:
    """Gate for platform-owner/admin-only endpoints (user management, manual
    subscription grants). Admins are flagged via users.role = 'admin'."""
    if current_user.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")
    return current_user

@app.get("/admin/demo-requests")
def admin_list_demo_requests(current_user: dict = Depends(require_admin)):
    conn = get_conn()
    rows = conn.execute("SELECT * FROM demo_requests ORDER BY id DESC").fetchall()
    conn.close()
    return [dict(r) for r in rows]

class DemoRequestStatusUpdate(BaseModel):
    status: Literal["new", "contacted", "converted", "declined"]

@app.post("/admin/demo-requests/{request_id}/status")
def admin_update_demo_request_status(request_id: int, payload: DemoRequestStatusUpdate, current_user: dict = Depends(require_admin)):
    conn = get_conn()
    existing = conn.execute("SELECT id FROM demo_requests WHERE id = ?", (request_id,)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Demo request not found")
    conn.execute("UPDATE demo_requests SET status = ? WHERE id = ?", (payload.status, request_id))
    conn.commit()
    conn.close()
    return {"id": request_id, "status": payload.status}

def get_subscription_row(conn, user_id: int) -> dict | None:
    row = conn.execute(
        "SELECT * FROM subscriptions WHERE user_id = ? ORDER BY id DESC LIMIT 1",
        (user_id,)
    ).fetchone()
    return dict(row) if row else None

def subscription_is_active(sub: dict | None) -> bool:
    if not sub:
        return False
    if sub["status"] not in ("trialing", "active"):
        return False
    if sub["current_period_end"]:
        try:
            if datetime.strptime(sub["current_period_end"], "%Y-%m-%d") < datetime.now():
                return False
        except ValueError:
            pass
    return True

def require_subscription(current_user: dict = Depends(get_current_user)) -> dict:
    """Gate for paid market-intelligence endpoints (currently: /opportunities).
    Admin accounts always pass -- they don't need a subscription to manage the
    platform they own. Everyone else needs an active trial/paid subscription,
    granted either manually by an admin (POST /admin/subscriptions/{id}/grant)
    or, once configured, via Stripe checkout (POST /subscription/checkout)."""
    if current_user.get("role") == "admin":
        return current_user
    conn = get_conn()
    sub = get_subscription_row(conn, current_user["id"])
    conn.close()
    if not subscription_is_active(sub):
        raise HTTPException(
            status_code=402,
            detail="An active subscription is required to view this section. "
                   "See /subscription/plans, or contact us to get started."
        )
    return current_user

# ---------------- API keys: programmatic access for Enterprise/Pro customers ----------------

def generate_api_key() -> tuple[str, str, str]:
    """Returns (raw_key_to_show_once, key_hash_to_store, key_prefix_for_display)."""
    raw = "mfgcc_" + secrets.token_urlsafe(32)
    key_hash = hashlib.sha256(raw.encode()).hexdigest()
    prefix = raw[:12] + "..."
    return raw, key_hash, prefix

def get_user_by_api_key(conn, raw_key: str) -> dict | None:
    key_hash = hashlib.sha256(raw_key.encode()).hexdigest()
    row = conn.execute(
        "SELECT * FROM api_keys WHERE key_hash = ? AND is_active = 1", (key_hash,)
    ).fetchone()
    if not row:
        return None
    conn.execute("UPDATE api_keys SET last_used_at = CURRENT_TIMESTAMP WHERE id = ?", (row["id"],))
    conn.commit()
    user = conn.execute("SELECT * FROM users WHERE id = ? AND is_active = 1", (row["user_id"],)).fetchone()
    return dict(user) if user else None

def require_subscription_flexible(
    authorization: str | None = Header(default=None),
    x_api_key: str | None = Header(default=None),
) -> dict:
    """Same gating as require_subscription, but accepts either a JWT session
    (Authorization: Bearer ...) or an API key (X-API-Key: ...) -- the latter
    is how Enterprise customers hit export/data endpoints programmatically
    without a browser session."""
    conn = get_conn()
    current_user = None
    if x_api_key:
        current_user = get_user_by_api_key(conn, x_api_key)
        if not current_user:
            conn.close()
            raise HTTPException(status_code=401, detail="Invalid or revoked API key")
    elif authorization and authorization.startswith("Bearer "):
        token = authorization.split(" ", 1)[1]
        try:
            payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        except jwt.ExpiredSignatureError:
            conn.close()
            raise HTTPException(status_code=401, detail="Session expired, please log in again")
        except jwt.InvalidTokenError:
            conn.close()
            raise HTTPException(status_code=401, detail="Invalid session token")
        user = conn.execute("SELECT * FROM users WHERE id = ? AND is_active = 1", (payload["sub"],)).fetchone()
        current_user = dict(user) if user else None
    if not current_user:
        conn.close()
        raise HTTPException(status_code=401, detail="Missing Authorization header or X-API-Key")
    if current_user.get("role") == "admin":
        conn.close()
        return current_user
    sub = get_subscription_row(conn, current_user["id"])
    conn.close()
    if not subscription_is_active(sub):
        raise HTTPException(
            status_code=402,
            detail="An active subscription is required to use this endpoint. "
                   "See /subscription/plans, or contact us to get started."
        )
    return current_user

class ApiKeyCreate(BaseModel):
    label: str = Field(..., min_length=1, max_length=100)

@app.post("/api-keys")
def create_api_key(payload: ApiKeyCreate, current_user: dict = Depends(require_subscription)):
    """Generates a new API key for the logged-in (subscribed) user. The raw
    key is returned exactly once -- only its hash is stored, matching how
    passwords are handled."""
    raw_key, key_hash, prefix = generate_api_key()
    conn = get_conn()
    cur = conn.execute(
        "INSERT INTO api_keys (user_id, key_hash, key_prefix, label) VALUES (?,?,?,?)",
        (current_user["id"], key_hash, prefix, payload.label)
    )
    conn.commit()
    key_id = cur.lastrowid
    conn.close()
    return {"id": key_id, "label": payload.label, "api_key": raw_key,
            "warning": "Save this key now -- it will not be shown again."}

@app.get("/api-keys")
def list_api_keys(current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    rows = conn.execute(
        "SELECT id, key_prefix, label, created_at, last_used_at, is_active FROM api_keys WHERE user_id = ? ORDER BY id DESC",
        (current_user["id"],)
    ).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.delete("/api-keys/{key_id}")
def revoke_api_key(key_id: int, current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    row = conn.execute("SELECT id FROM api_keys WHERE id = ? AND user_id = ?", (key_id, current_user["id"])).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="API key not found")
    conn.execute("UPDATE api_keys SET is_active = 0 WHERE id = ?", (key_id,))
    conn.commit()
    conn.close()
    return {"id": key_id, "is_active": False}

# ---------------- Data export (CSV): Pro/Enterprise only, JWT or API key ----------------

def _csv_response(rows: list[dict], filename: str) -> Response:
    if not rows:
        return Response(content="", media_type="text/csv", headers={"Content-Disposition": f"attachment; filename={filename}"})
    buf = io.StringIO()
    writer = csv.DictWriter(buf, fieldnames=list(rows[0].keys()))
    writer.writeheader()
    writer.writerows(rows)
    return Response(content=buf.getvalue(), media_type="text/csv",
                     headers={"Content-Disposition": f"attachment; filename={filename}"})

@app.get("/export/companies.csv")
def export_companies_csv(current_user: dict = Depends(require_subscription_flexible)):
    conn = get_conn()
    rows = conn.execute("SELECT name, category, origin, headquarters, website, ksa_status, status_tag FROM manufacturers WHERE is_published = 1 ORDER BY name").fetchall()
    conn.close()
    return _csv_response([dict(r) for r in rows], "medforsa-companies.csv")

@app.get("/export/products.csv")
def export_products_csv(current_user: dict = Depends(require_subscription_flexible)):
    conn = get_conn()
    rows = conn.execute("""
        SELECT p.product_name, p.product_type, m.name as manufacturer, p.department, p.throughput
        FROM products p LEFT JOIN manufacturers m ON p.manufacturer_id = m.id ORDER BY m.name, p.product_name
    """).fetchall()
    conn.close()
    return _csv_response([dict(r) for r in rows], "medforsa-products.csv")

@app.get("/export/distributors.csv")
def export_distributors_csv(current_user: dict = Depends(require_subscription_flexible)):
    conn = get_conn()
    rows = conn.execute("SELECT name, country, represents, market_strength_tier FROM distributors ORDER BY name").fetchall()
    conn.close()
    return _csv_response([dict(r) for r in rows], "medforsa-distributors.csv")

@app.get("/export/opportunities.csv")
def export_opportunities_csv(current_user: dict = Depends(require_subscription_flexible)):
    conn = get_conn()
    rows = conn.execute("""
        SELECT m.name as company_name, o.reason, o.action,
               (o.score_no_distributor + o.score_confidence + o.score_brand) as total_score
        FROM opportunities o JOIN manufacturers m ON m.id = o.manufacturer_id ORDER BY total_score DESC
    """).fetchall()
    conn.close()
    return _csv_response([dict(r) for r in rows], "medforsa-opportunities.csv")

@app.get("/export/lab-tests.csv")
def export_lab_tests_csv(current_user: dict = Depends(require_subscription_flexible)):
    conn = get_conn()
    rows = conn.execute("SELECT name_en, category, specimen_type, purpose_en FROM lab_tests WHERE is_published = 1 ORDER BY category, name_en").fetchall()
    conn.close()
    return _csv_response([dict(r) for r in rows], "medforsa-lab-tests.csv")

class RegisterRequest(BaseModel):
    email: EmailStr
    password: str = Field(..., min_length=8)
    company_name: str = Field(..., min_length=1)
    full_name: str | None = None

class LoginRequest(BaseModel):
    email: EmailStr
    password: str

@app.post("/auth/register")
def register(reg: RegisterRequest):
    conn = get_conn()
    existing = conn.execute("SELECT id FROM users WHERE email = ?", (reg.email,)).fetchone()
    if existing:
        conn.close()
        raise HTTPException(status_code=400, detail="An account with this email already exists")
    password_hash = bcrypt.hashpw(reg.password.encode(), bcrypt.gensalt()).decode()
    cur = conn.execute(
        "INSERT INTO users (email, password_hash, company_name, full_name, role) VALUES (?,?,?,?,'user')",
        (reg.email, password_hash, reg.company_name, reg.full_name)
    )
    conn.commit()
    user_id = cur.lastrowid
    conn.execute(
        "INSERT INTO subscriptions (user_id, plan, status, current_period_end, granted_by, grant_note) VALUES (?,?,?,?,?,?)",
        (user_id, "trial", "trialing", (datetime.now() + timedelta(days=14)).strftime("%Y-%m-%d"),
         "system", "Automatic 14-day trial on signup")
    )
    conn.execute("INSERT INTO audit_log (table_name, record_id, action, detail) VALUES ('users', ?, 'insert', ?)",
                 (user_id, f"New account registered: {reg.email} ({reg.company_name})"))
    conn.commit()
    conn.close()
    token = create_access_token(user_id, reg.email)
    return {"access_token": token, "token_type": "bearer",
            "user": {"id": user_id, "email": reg.email, "company_name": reg.company_name, "full_name": reg.full_name, "role": "user"}}

@app.post("/auth/login")
def login(creds: LoginRequest):
    conn = get_conn()
    user = conn.execute("SELECT * FROM users WHERE email = ? AND is_active = 1", (creds.email,)).fetchone()
    conn.close()
    if not user or not bcrypt.checkpw(creds.password.encode(), user["password_hash"].encode()):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    token = create_access_token(user["id"], user["email"])
    return {"access_token": token, "token_type": "bearer", "user": public_user_fields(dict(user))}

def send_email(to_email: str, subject: str, body: str) -> bool:
    """Generic SMTP sender -- returns True if actually sent, False if SMTP
    isn't configured yet (SMTP_HOST/USER/PASSWORD env vars, same ready-but-
    needs-credentials pattern as Stripe/Twilio)."""
    smtp_host = os.environ.get("SMTP_HOST")
    smtp_user = os.environ.get("SMTP_USER")
    smtp_password = os.environ.get("SMTP_PASSWORD")
    smtp_from = os.environ.get("SMTP_FROM", smtp_user or "no-reply@medforsagcc.com")
    if not (smtp_host and smtp_user and smtp_password):
        return False
    smtp_port = int(os.environ.get("SMTP_PORT", "587"))
    msg = MIMEText(body)
    msg["Subject"] = subject
    msg["From"] = smtp_from
    msg["To"] = to_email
    with smtplib.SMTP(smtp_host, smtp_port) as server:
        server.starttls()
        server.login(smtp_user, smtp_password)
        server.sendmail(smtp_from, [to_email], msg.as_string())
    return True

def send_password_reset_email(to_email: str, code: str) -> bool:
    """Sends the reset code via SMTP if credentials are configured (same
    pattern as Stripe/Twilio: code is ready, just needs env vars). Returns
    True if an email was actually sent, False if SMTP isn't configured yet
    (caller falls back to logging the code server-side)."""
    return send_email(
        to_email,
        "MedForsa GCC -- password reset code",
        f"Your MedForsa GCC password reset code is: {code}\n\n"
        f"This code expires in 15 minutes. If you didn't request this, you can ignore this email."
    )

class ForgotPasswordRequest(BaseModel):
    email: EmailStr

@app.post("/auth/forgot-password")
def forgot_password(payload: ForgotPasswordRequest):
    """Enter the email you registered with -- if it matches an account, a
    6-digit verification code is generated (emailed if SMTP is configured,
    otherwise logged server-side so the account owner/admin can retrieve it
    until email sending is set up). Always returns the same generic message
    regardless of whether the email exists, so this can't be used to probe
    which emails are registered."""
    conn = get_conn()
    user = conn.execute("SELECT id, email FROM users WHERE email = ? AND is_active = 1", (payload.email,)).fetchone()
    if user:
        code = f"{secrets.randbelow(1000000):06d}"
        expires_at = (datetime.now() + timedelta(minutes=15)).isoformat()
        conn.execute(
            "INSERT INTO password_resets (user_id, code, expires_at) VALUES (?, ?, ?)",
            (user["id"], code, expires_at)
        )
        conn.commit()
        emailed = send_password_reset_email(user["email"], code)
        if not emailed:
            logger.info(f"PASSWORD RESET CODE (SMTP not configured, logging instead) for {user['email']}: {code} (expires in 15 min)")
    conn.close()
    return {"message": "If that email is registered, a verification code has been sent."}

class ResetPasswordRequest(BaseModel):
    email: EmailStr
    code: str = Field(..., min_length=6, max_length=6)
    new_password: str = Field(..., min_length=8)

@app.post("/auth/reset-password")
def reset_password(payload: ResetPasswordRequest):
    conn = get_conn()
    user = conn.execute("SELECT id FROM users WHERE email = ? AND is_active = 1", (payload.email,)).fetchone()
    if not user:
        conn.close()
        raise HTTPException(status_code=400, detail="Invalid code or email")
    reset_row = conn.execute(
        "SELECT id, expires_at FROM password_resets WHERE user_id = ? AND code = ? AND used = 0 ORDER BY id DESC LIMIT 1",
        (user["id"], payload.code)
    ).fetchone()
    if not reset_row:
        conn.close()
        raise HTTPException(status_code=400, detail="Invalid code or email")
    if datetime.fromisoformat(reset_row["expires_at"]) < datetime.now():
        conn.close()
        raise HTTPException(status_code=400, detail="This code has expired -- request a new one")
    new_hash = bcrypt.hashpw(payload.new_password.encode(), bcrypt.gensalt()).decode()
    conn.execute("UPDATE users SET password_hash = ? WHERE id = ?", (new_hash, user["id"]))
    conn.execute("UPDATE password_resets SET used = 1 WHERE id = ?", (reset_row["id"],))
    conn.execute("INSERT INTO audit_log (table_name, record_id, action, detail) VALUES ('users', ?, 'update', ?)",
                 (user["id"], "Password reset via forgot-password flow"))
    conn.commit()
    conn.close()
    return {"message": "Password updated -- you can now log in with your new password."}

@app.get("/admin/password-resets/pending")
def admin_pending_password_resets(current_user: dict = Depends(require_admin)):
    """Until SMTP is configured, this lets an admin see recently-generated
    reset codes (e.g. their own) instead of digging through server logs."""
    conn = get_conn()
    rows = conn.execute("""
        SELECT pr.id, u.email, pr.code, pr.expires_at, pr.used, pr.created_at
        FROM password_resets pr JOIN users u ON u.id = pr.user_id
        WHERE pr.created_at >= datetime('now', '-1 day')
        ORDER BY pr.id DESC
    """).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/auth/me")
def read_me(current_user: dict = Depends(get_current_user)):
    return public_user_fields(current_user)

# ---------------- Subscription plans, status, and Stripe checkout ----------------

SUBSCRIPTION_PLANS = {
    "trial": {"name": "Trial", "price_sar": 0, "billing": "14 days, one-time",
               "features": ["Full platform access for 14 days", "No card required"]},
    "pro": {"name": "Pro", "price_sar": 1500, "billing": "per month",
             "features": ["Opportunities & market intelligence", "Distributor & regulatory data",
                          "Doctor AI", "Unlimited leads & watchlist"]},
    "enterprise": {"name": "Enterprise", "price_sar": None, "billing": "annual, contact us",
                    "features": ["Everything in Pro", "Multi-user team accounts",
                                 "API access", "Priority data requests"]},
}

@app.get("/subscription/plans")
def subscription_plans():
    return SUBSCRIPTION_PLANS

@app.get("/subscription/status")
def subscription_status(current_user: dict = Depends(get_current_user)):
    if current_user.get("role") == "admin":
        return {"plan": "admin", "status": "active", "is_active": True, "current_period_end": None}
    conn = get_conn()
    sub = get_subscription_row(conn, current_user["id"])
    conn.close()
    if not sub:
        return {"plan": None, "status": "none", "is_active": False, "current_period_end": None}
    return {"plan": sub["plan"], "status": sub["status"], "is_active": subscription_is_active(sub),
            "current_period_end": sub["current_period_end"]}

class CheckoutRequest(BaseModel):
    plan: Literal["pro", "enterprise"]

@app.post("/subscription/checkout")
def subscription_checkout(payload: CheckoutRequest, current_user: dict = Depends(get_current_user)):
    """Creates a Stripe Checkout session for self-serve payment. Requires
    STRIPE_SECRET_KEY and a STRIPE_PRICE_ID_<PLAN> env var to be set on the
    server (same pattern as the WhatsApp/Twilio integration: backend is
    ready, just needs real credentials filled in before it can process a
    real payment). Until then, this returns a clear 501 rather than pretending
    to charge a card."""
    stripe_key = os.environ.get("STRIPE_SECRET_KEY")
    price_id = os.environ.get(f"STRIPE_PRICE_ID_{payload.plan.upper()}")
    if not stripe_key or not price_id:
        raise HTTPException(
            status_code=501,
            detail="Online payment isn't configured yet. Please contact us directly to "
                   "activate a subscription, or ask an admin to grant one."
        )
    import stripe
    stripe.api_key = stripe_key
    site_url = os.environ.get("SITE_URL", "https://medintel-gcc.onrender.com")
    session = stripe.checkout.Session.create(
        mode="subscription",
        customer_email=current_user["email"],
        line_items=[{"price": price_id, "quantity": 1}],
        success_url=f"{site_url}/app?subscription=success",
        cancel_url=f"{site_url}/app?subscription=cancelled",
        metadata={"user_id": str(current_user["id"]), "plan": payload.plan},
    )
    return {"checkout_url": session.url}

@app.post("/subscription/webhook")
async def subscription_webhook(request: Request):
    """Stripe webhook: activates a subscription in our DB once Stripe confirms
    payment. Requires STRIPE_WEBHOOK_SECRET to be set for signature verification."""
    webhook_secret = os.environ.get("STRIPE_WEBHOOK_SECRET")
    if not webhook_secret:
        raise HTTPException(status_code=501, detail="Stripe webhook not configured")
    import stripe
    payload = await request.body()
    sig_header = request.headers.get("stripe-signature", "")
    try:
        event = stripe.Webhook.construct_event(payload, sig_header, webhook_secret)
    except (ValueError, Exception):
        raise HTTPException(status_code=400, detail="Invalid webhook signature")
    if event["type"] == "checkout.session.completed":
        session = event["data"]["object"]
        user_id = int(session["metadata"]["user_id"])
        plan = session["metadata"]["plan"]
        conn = get_conn()
        conn.execute(
            "INSERT INTO subscriptions (user_id, plan, status, current_period_end, stripe_customer_id, stripe_subscription_id) "
            "VALUES (?,?,?,?,?,?)",
            (user_id, plan, "active", (datetime.now() + timedelta(days=365)).strftime("%Y-%m-%d"),
             session.get("customer"), session.get("subscription"))
        )
        conn.commit()
        conn.close()
    return {"received": True}

# ---------------- Admin: user & subscription management ----------------

@app.get("/admin/users")
def admin_list_users(current_user: dict = Depends(require_admin)):
    conn = get_conn()
    users = conn.execute("SELECT id, email, company_name, full_name, role, is_active, created_at FROM users ORDER BY id").fetchall()
    result = []
    for u in users:
        u = dict(u)
        sub = get_subscription_row(conn, u["id"])
        u["subscription"] = {"plan": sub["plan"], "status": sub["status"], "is_active": subscription_is_active(sub),
                              "current_period_end": sub["current_period_end"]} if sub else None
        result.append(u)
    conn.close()
    return result

@app.get("/admin/stats")
def admin_stats(current_user: dict = Depends(require_admin)):
    conn = get_conn()
    total_users = conn.execute("SELECT COUNT(*) c FROM users").fetchone()["c"]
    active_subs = conn.execute(
        "SELECT COUNT(DISTINCT user_id) c FROM subscriptions WHERE status IN ('trialing','active')"
    ).fetchone()["c"]
    total_leads = conn.execute("SELECT COUNT(*) c FROM leads").fetchone()["c"]
    conn.close()
    return {"total_users": total_users, "active_subscriptions": active_subs, "total_leads": total_leads}

class RoleUpdate(BaseModel):
    role: Literal["user", "admin"]

@app.post("/admin/users/{user_id}/role")
def admin_set_role(user_id: int, payload: RoleUpdate, current_user: dict = Depends(require_admin)):
    conn = get_conn()
    existing = conn.execute("SELECT id, email FROM users WHERE id = ?", (user_id,)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="User not found")
    conn.execute("UPDATE users SET role = ? WHERE id = ?", (payload.role, user_id))
    conn.execute("INSERT INTO audit_log (table_name, record_id, action, detail) VALUES ('users', ?, 'update', ?)",
                 (user_id, f"Role changed to '{payload.role}' by admin {current_user['email']}"))
    conn.commit()
    conn.close()
    return {"id": user_id, "role": payload.role}

class ActiveUpdate(BaseModel):
    is_active: bool

@app.post("/admin/users/{user_id}/active")
def admin_set_active(user_id: int, payload: ActiveUpdate, current_user: dict = Depends(require_admin)):
    conn = get_conn()
    existing = conn.execute("SELECT id, email FROM users WHERE id = ?", (user_id,)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="User not found")
    conn.execute("UPDATE users SET is_active = ? WHERE id = ?", (int(payload.is_active), user_id))
    conn.execute("INSERT INTO audit_log (table_name, record_id, action, detail) VALUES ('users', ?, 'update', ?)",
                 (user_id, f"Account {'enabled' if payload.is_active else 'disabled'} by admin {current_user['email']}"))
    conn.commit()
    conn.close()
    return {"id": user_id, "is_active": payload.is_active}

class SubscriptionGrant(BaseModel):
    plan: Literal["trial", "pro", "enterprise"]
    days: int = Field(365, ge=1, le=3650)
    note: str | None = None

@app.post("/admin/subscriptions/{user_id}/grant")
def admin_grant_subscription(user_id: int, payload: SubscriptionGrant, current_user: dict = Depends(require_admin)):
    """Manually activate/extend a subscription -- the primary path for KSA/GCC
    B2B deals closed offline (bank transfer, PO, invoice) rather than through
    self-serve card checkout."""
    conn = get_conn()
    existing = conn.execute("SELECT id, email FROM users WHERE id = ?", (user_id,)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="User not found")
    period_end = (datetime.now() + timedelta(days=payload.days)).strftime("%Y-%m-%d")
    conn.execute(
        "INSERT INTO subscriptions (user_id, plan, status, current_period_end, granted_by, grant_note) VALUES (?,?,?,?,?,?)",
        (user_id, payload.plan, "active", period_end, current_user["email"], payload.note)
    )
    conn.execute("INSERT INTO audit_log (table_name, record_id, action, detail) VALUES ('users', ?, 'update', ?)",
                 (user_id, f"Subscription '{payload.plan}' granted through {period_end} by admin {current_user['email']}"))
    conn.commit()
    conn.close()
    return {"id": user_id, "plan": payload.plan, "status": "active", "current_period_end": period_end}

@app.post("/admin/subscriptions/{user_id}/revoke")
def admin_revoke_subscription(user_id: int, current_user: dict = Depends(require_admin)):
    conn = get_conn()
    sub = get_subscription_row(conn, user_id)
    if not sub:
        conn.close()
        raise HTTPException(status_code=404, detail="No subscription found for this user")
    conn.execute("UPDATE subscriptions SET status = 'canceled', updated_at = CURRENT_TIMESTAMP WHERE id = ?", (sub["id"],))
    conn.execute("INSERT INTO audit_log (table_name, record_id, action, detail) VALUES ('users', ?, 'update', ?)",
                 (user_id, f"Subscription canceled by admin {current_user['email']}"))
    conn.commit()
    conn.close()
    return {"id": user_id, "status": "canceled"}

@app.get("/opportunities")
def list_opportunities():
    """TEMPORARY: opened without login/subscription while the platform is
    still being actively built out, per explicit request (2026-07-21).
    Before real launch, restore `current_user: dict = Depends(require_subscription)`
    here to re-gate this behind login + an active subscription -- it reveals
    which manufacturers have no confirmed KSA distributor, sensitive
    competitive strategy info that shouldn't stay open to anonymous visitors
    once this becomes a real commercial product."""
    conn = get_conn()
    rows = conn.execute(
        """SELECT o.*, m.name as company_name FROM opportunities o
           JOIN manufacturers m ON m.id = o.manufacturer_id"""
    ).fetchall()
    conn.close()
    results = []
    for r in rows:
        d = dict(r)
        d["total_score"] = d["score_no_distributor"] + d["score_confidence"] + d["score_brand"]
        results.append(d)
    return sorted(results, key=lambda x: -x["total_score"])

class NewLead(BaseModel):
    manufacturer_id: int
    contact_name: str | None = None
    contact_role: str | None = None
    contact_email: str | None = None
    contact_phone: str | None = None
    status: Literal["new", "contacted", "in_progress", "won", "lost"] = "new"
    notes: str | None = None

class UpdateLead(BaseModel):
    contact_name: str | None = None
    contact_role: str | None = None
    contact_email: str | None = None
    contact_phone: str | None = None
    status: Literal["new", "contacted", "in_progress", "won", "lost"] | None = None
    notes: str | None = None

class NewInteraction(BaseModel):
    interaction_type: Literal["call", "email", "meeting", "note"]
    summary: str = Field(..., min_length=1)
    next_followup_date: str | None = None

@app.get("/leads")
def list_leads(current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    rows = conn.execute("""
        SELECT l.*, m.name as company_name, m.category, m.status_tag
        FROM leads l JOIN manufacturers m ON m.id = l.manufacturer_id
        WHERE l.user_id = ?
        ORDER BY l.created_at DESC
    """, (current_user["id"],)).fetchall()
    leads = [dict(r) for r in rows]
    for lead in leads:
        interactions = conn.execute(
            "SELECT * FROM lead_interactions WHERE lead_id = ? ORDER BY interaction_date DESC", (lead["id"],)
        ).fetchall()
        lead["interactions"] = [dict(i) for i in interactions]
        lead["last_interaction"] = lead["interactions"][0] if lead["interactions"] else None
    conn.close()
    return leads

@app.post("/leads")
def create_lead(lead: NewLead, current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    company = conn.execute("SELECT id FROM manufacturers WHERE id = ?", (lead.manufacturer_id,)).fetchone()
    if not company:
        conn.close()
        raise HTTPException(status_code=404, detail="manufacturer_id not found")
    cur = conn.execute(
        "INSERT INTO leads (user_id, manufacturer_id, contact_name, contact_role, contact_email, contact_phone, status, notes) VALUES (?,?,?,?,?,?,?,?)",
        (current_user["id"], lead.manufacturer_id, lead.contact_name, lead.contact_role, lead.contact_email, lead.contact_phone, lead.status, lead.notes)
    )
    conn.commit()
    lead_id = cur.lastrowid
    conn.execute("INSERT INTO audit_log (table_name, record_id, action, detail) VALUES ('leads', ?, 'insert', ?)",
                 (lead_id, f"New lead created by user {current_user['id']} for manufacturer_id {lead.manufacturer_id}"))
    conn.commit()
    conn.close()
    return {"id": lead_id, "status": "created"}

def _owned_lead_or_404(conn, lead_id: int, user_id: int):
    row = conn.execute("SELECT * FROM leads WHERE id = ? AND user_id = ?", (lead_id, user_id)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="Lead not found")
    return row

@app.put("/leads/{lead_id}")
def update_lead(lead_id: int, updates: UpdateLead, current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    _owned_lead_or_404(conn, lead_id, current_user["id"])
    fields = {k: v for k, v in updates.dict().items() if v is not None}
    if fields:
        set_clause = ", ".join(f"{k} = ?" for k in fields)
        conn.execute(f"UPDATE leads SET {set_clause} WHERE id = ?", (*fields.values(), lead_id))
        conn.commit()
    conn.close()
    return {"id": lead_id, "status": "updated"}

@app.delete("/leads/{lead_id}")
def delete_lead(lead_id: int, current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    _owned_lead_or_404(conn, lead_id, current_user["id"])
    conn.execute("DELETE FROM leads WHERE id = ?", (lead_id,))
    conn.commit()
    conn.close()
    return {"id": lead_id, "status": "deleted"}

@app.post("/leads/{lead_id}/interactions")
def add_interaction(lead_id: int, interaction: NewInteraction, current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    _owned_lead_or_404(conn, lead_id, current_user["id"])
    conn.execute(
        "INSERT INTO lead_interactions (lead_id, interaction_type, summary, next_followup_date) VALUES (?,?,?,?)",
        (lead_id, interaction.interaction_type, interaction.summary, interaction.next_followup_date)
    )
    conn.commit()
    conn.close()
    return {"lead_id": lead_id, "status": "interaction logged"}

@app.get("/watchlist")
def list_watchlist(current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    rows = conn.execute("""
        SELECT w.id as watchlist_id, w.added_at, m.* FROM watchlist w
        JOIN manufacturers m ON m.id = w.manufacturer_id
        WHERE w.user_id = ?
        ORDER BY w.added_at DESC
    """, (current_user["id"],)).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.post("/watchlist/{manufacturer_id}")
def add_to_watchlist(manufacturer_id: int, current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    company = conn.execute("SELECT id FROM manufacturers WHERE id = ?", (manufacturer_id,)).fetchone()
    if not company:
        conn.close()
        raise HTTPException(status_code=404, detail="manufacturer_id not found")
    try:
        conn.execute("INSERT INTO watchlist (user_id, manufacturer_id) VALUES (?,?)", (current_user["id"], manufacturer_id))
        conn.commit()
    except sqlite3.IntegrityError:
        pass
    conn.close()
    return {"manufacturer_id": manufacturer_id, "status": "watched"}

@app.delete("/watchlist/{manufacturer_id}")
def remove_from_watchlist(manufacturer_id: int, current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    conn.execute("DELETE FROM watchlist WHERE manufacturer_id = ? AND user_id = ?", (manufacturer_id, current_user["id"]))
    conn.commit()
    conn.close()
    return {"manufacturer_id": manufacturer_id, "status": "unwatched"}

@app.get("/daily-briefing")
def daily_briefing(current_user: dict = Depends(get_current_user)):
    """Executive daily briefing: what's genuinely new since recent activity --
    derived entirely from real audit_log/opportunities/conferences/leads data,
    scoped to the logged-in account's own leads/watchlist. Nothing generated
    or invented."""
    conn = get_conn()
    recent_changes = conn.execute("SELECT * FROM audit_log ORDER BY id DESC LIMIT 8").fetchall()
    top_opps = conn.execute("""
        SELECT o.*, m.name as company_name FROM opportunities o
        JOIN manufacturers m ON m.id = o.manufacturer_id
    """).fetchall()
    top_opps_scored = sorted(
        [dict(r, total_score=r["score_no_distributor"] + r["score_confidence"] + r["score_brand"]) for r in top_opps],
        key=lambda x: -x["total_score"]
    )[:3]
    upcoming = conn.execute("SELECT * FROM conferences ORDER BY event_date").fetchall()
    today = datetime.now()
    upcoming_soon = [dict(c) for c in upcoming if 0 <= (datetime.strptime(c["event_date"], "%Y-%m-%d") - today).days <= 30]
    stale_leads = conn.execute("""
        SELECT l.*, m.name as company_name FROM leads l
        JOIN manufacturers m ON m.id = l.manufacturer_id
        WHERE l.user_id = ? AND l.status NOT IN ('won','lost')
        AND l.id NOT IN (
            SELECT lead_id FROM lead_interactions
            WHERE interaction_date >= date('now', '-14 days')
        )
    """, (current_user["id"],)).fetchall()
    watchlist_count = conn.execute("SELECT count(*) c FROM watchlist WHERE user_id = ?", (current_user["id"],)).fetchone()["c"]
    conn.close()
    return {
        "generated_at": today.strftime("%Y-%m-%d %H:%M"),
        "recent_changes": [dict(r) for r in recent_changes],
        "top_opportunities": top_opps_scored,
        "conferences_next_30_days": upcoming_soon,
        "stale_leads_needing_followup": [dict(r) for r in stale_leads],
        "watchlist_count": watchlist_count,
        "note": "Every field above is computed live from real data -- nothing generated or invented.",
    }

# ---------------- Notifications: computed fresh from real data, no cron job needed ----------------
# A notification is never stored as its own row -- it's derived live from opportunities,
# conferences, and company_distributors (each has created_at now), then cross-referenced
# against notification_reads to know what THIS user has already seen. This avoids any
# risk of stale/duplicated notification content drifting from the real underlying data.

def _build_notification_candidates(conn):
    candidates = []

    new_opps = conn.execute("""
        SELECT o.id, o.created_at, m.name as company_name, o.reason
        FROM opportunities o JOIN manufacturers m ON m.id = o.manufacturer_id
        WHERE o.created_at >= date('now', '-30 days')
        ORDER BY o.created_at DESC
    """).fetchall()
    for r in new_opps:
        candidates.append({
            "key": f"opportunity:{r['id']}", "type": "opportunity", "created_at": r["created_at"],
            "title": f"New opportunity: {r['company_name']}", "detail": r["reason"],
            "url": "#opportunities",
        })

    upcoming_conf = conn.execute("""
        SELECT id, name, event_date, place FROM conferences
        WHERE event_date >= date('now') AND event_date <= date('now', '+30 days')
        ORDER BY event_date ASC
    """).fetchall()
    for r in upcoming_conf:
        candidates.append({
            "key": f"conference-soon:{r['id']}", "type": "conference",
            "created_at": r["event_date"],  # ranked by proximity, not insert date
            "title": f"{r['name']} is coming up", "detail": f"{r['event_date']} -- {r['place']}",
            "url": "#conferences",
        })

    new_conf = conn.execute("""
        SELECT id, name, event_date, place FROM conferences
        WHERE created_at >= date('now', '-30 days')
        ORDER BY created_at DESC
    """).fetchall()
    for r in new_conf:
        candidates.append({
            "key": f"conference-new:{r['id']}", "type": "conference",
            "created_at": r["created_at"],
            "title": f"New conference added: {r['name']}", "detail": f"{r['event_date']} -- {r['place']}",
            "url": "#conferences",
        })

    new_links = conn.execute("""
        SELECT cd.id, cd.created_at, m.name as company_name, d.name as distributor_name
        FROM company_distributors cd
        JOIN manufacturers m ON m.id = cd.manufacturer_id
        JOIN distributors d ON d.id = cd.distributor_id
        WHERE cd.created_at >= date('now', '-30 days')
        ORDER BY cd.created_at DESC
    """).fetchall()
    for r in new_links:
        candidates.append({
            "key": f"distributor-link:{r['id']}", "type": "distributor_change",
            "created_at": r["created_at"],
            "title": f"Distributor confirmed: {r['company_name']}",
            "detail": f"Now linked to {r['distributor_name']}", "url": "#distributors",
        })

    candidates.sort(key=lambda c: c["created_at"], reverse=True)
    return candidates

def _compose_digest_email(user_email: str, candidates: list) -> str:
    lines = [
        f"Your MedForsa GCC update digest ({len(candidates)} new item{'s' if len(candidates) != 1 else ''}):",
        "",
    ]
    for c in candidates:
        lines.append(f"- {c['title']}")
        if c.get("detail"):
            lines.append(f"  {c['detail']}")
    lines.append("")
    lines.append(f"View them at {SITE_URL}/app")
    return "\n".join(lines)

@app.post("/admin/send-digest-emails")
def admin_send_digest_emails(current_user: dict = Depends(require_admin)):
    """Manually triggers a digest email to every subscribed user with unread
    notifications. No cron scheduler exists in this environment, so this is
    admin-triggered for now -- to automate it, point an external scheduler
    (e.g. a free cron-job.org ping, or a Render Cron Job on a paid plan) at
    this endpoint with the admin's API key on a daily/weekly schedule."""
    conn = get_conn()
    candidates = _build_notification_candidates(conn)
    if not candidates:
        conn.close()
        return {"sent": 0, "message": "No new items to send -- nothing is out of date."}
    users = conn.execute("SELECT id, email FROM users WHERE is_active = 1").fetchall()
    sent, skipped_no_smtp = 0, 0
    for u in users:
        read_keys = {r["notification_key"] for r in conn.execute(
            "SELECT notification_key FROM notification_reads WHERE user_id = ?", (u["id"],)
        ).fetchall()}
        unread = [c for c in candidates if c["key"] not in read_keys]
        if not unread:
            continue
        body = _compose_digest_email(u["email"], unread)
        emailed = send_email(u["email"], f"MedForsa GCC -- {len(unread)} new update(s)", body)
        if emailed:
            sent += 1
        else:
            skipped_no_smtp += 1
    conn.close()
    if skipped_no_smtp and not sent:
        return {"sent": 0, "message": f"SMTP not configured -- would have sent to {skipped_no_smtp} user(s). Set SMTP_HOST/SMTP_USER/SMTP_PASSWORD to enable."}
    return {"sent": sent, "message": f"Digest sent to {sent} user(s)."}

@app.get("/notifications")
def list_notifications(current_user: dict = Depends(get_current_user), include_read: bool = False, limit: int = 30):
    conn = get_conn()
    candidates = _build_notification_candidates(conn)
    read_keys = {r["notification_key"] for r in conn.execute(
        "SELECT notification_key FROM notification_reads WHERE user_id = ?", (current_user["id"],)
    ).fetchall()}
    conn.close()
    for c in candidates:
        c["is_read"] = c["key"] in read_keys
    if not include_read:
        candidates = [c for c in candidates if not c["is_read"]]
    return candidates[:limit]

@app.get("/notifications/unread-count")
def notifications_unread_count(current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    candidates = _build_notification_candidates(conn)
    read_keys = {r["notification_key"] for r in conn.execute(
        "SELECT notification_key FROM notification_reads WHERE user_id = ?", (current_user["id"],)
    ).fetchall()}
    conn.close()
    unread = sum(1 for c in candidates if c["key"] not in read_keys)
    return {"unread_count": unread}

class NotificationReadRequest(BaseModel):
    key: str

@app.post("/notifications/read")
def mark_notification_read(payload: NotificationReadRequest, current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    try:
        conn.execute(
            "INSERT OR IGNORE INTO notification_reads (user_id, notification_key) VALUES (?, ?)",
            (current_user["id"], payload.key)
        )
        conn.commit()
    finally:
        conn.close()
    return {"key": payload.key, "is_read": True}

@app.post("/notifications/read-all")
def mark_all_notifications_read(current_user: dict = Depends(get_current_user)):
    conn = get_conn()
    candidates = _build_notification_candidates(conn)
    for c in candidates:
        conn.execute(
            "INSERT OR IGNORE INTO notification_reads (user_id, notification_key) VALUES (?, ?)",
            (current_user["id"], c["key"])
        )
    conn.commit()
    conn.close()
    return {"marked_read": len(candidates)}

@app.get("/conferences")
def list_conferences(upcoming_only: bool = False):
    conn = get_conn()
    rows = conn.execute("SELECT * FROM conferences ORDER BY event_date").fetchall()
    conn.close()
    today = datetime.now()
    results = []
    for r in rows:
        d = dict(r)
        event_date = datetime.strptime(d["event_date"], "%Y-%m-%d")
        d["days_from_today"] = (event_date - today).days
        d["is_past"] = d["days_from_today"] < 0
        if upcoming_only and d["is_past"]:
            continue
        results.append(d)
    return results

@app.get("/distributors")
def list_distributors():
    conn = get_conn()
    rows = conn.execute("SELECT * FROM distributors").fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/dashboard")
def dashboard():
    conn = get_conn()
    total = conn.execute("SELECT count(*) c FROM manufacturers").fetchone()["c"]
    resolved = conn.execute("SELECT count(*) c FROM manufacturers WHERE status_tag IN ('covered','open')").fetchone()["c"]
    open_opps = conn.execute("SELECT count(*) c FROM manufacturers WHERE status_tag='open'").fetchone()["c"]
    conn.close()
    return {
        "companies_tracked": total,
        "open_opportunities": open_opps,
        "data_completeness_pct": round(resolved / total * 100) if total else 0,
        "snapshot_date": datetime.now().strftime("%Y-%m-%d"),
    }

@app.get("/audit-log")
def audit_log():
    conn = get_conn()
    rows = conn.execute("SELECT * FROM audit_log ORDER BY id DESC LIMIT 50").fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/technologies")
def list_technologies():
    conn = get_conn()
    techs = conn.execute("SELECT * FROM technologies").fetchall()
    result = []
    for t in techs:
        companies = conn.execute(
            """SELECT m.name FROM manufacturers m
               JOIN manufacturer_technology mt ON mt.manufacturer_id = m.id
               WHERE mt.technology_id = ?""", (t["id"],)
        ).fetchall()
        d = dict(t)
        d["companies"] = [c["name"] for c in companies]
        result.append(d)
    conn.close()
    return result

def generate_insight(company: dict) -> str:
    """Rule-based summary generator — plain templating over real fields, not a live AI call."""
    name = company["name"]
    parts = []
    if company["status_tag"] == "open":
        parts.append(f"no confirmed KSA distributor was found for {name} in our research")
    elif company["status_tag"] == "unclear":
        parts.append(f"{name}'s KSA distribution status is unconfirmed — worth a direct question")
    else:
        parts.append(f"{name} already has KSA coverage, so this is not an open distribution lead")

    if company.get("technologies"):
        techs = ", ".join(t["name"] for t in company["technologies"])
        parts.append(f"its confirmed focus areas are: {techs}")

    if company.get("opportunity"):
        total = company["opportunity"]["score_no_distributor"] + company["opportunity"]["score_confidence"] + company["opportunity"]["score_brand"]
        parts.append(f"it scores {total}/10 on our opportunity formula")

    parts.append(f"confidence tier: {company['confidence_tier']} ({company['sources'][:60]}...)" if len(company.get('sources',''))>60 else f"confidence tier: {company['confidence_tier']} ({company.get('sources','no sources yet')})")
    return ". ".join(p[0].upper()+p[1:] for p in parts) + "."

@app.get("/compare")
def compare_products(product_ids: str):
    """Comparison engine. Pass product_ids as a comma-separated list, e.g.
    /compare?product_ids=1,5,12. Returns each product's full record plus its
    manufacturer's origin/category and KSA distributor(s) for side-by-side display."""
    try:
        ids = [int(x) for x in product_ids.split(",") if x.strip()]
    except ValueError:
        raise HTTPException(status_code=400, detail="product_ids must be a comma-separated list of integers")
    if not ids:
        raise HTTPException(status_code=400, detail="At least one product_id is required")
    if len(ids) > 6:
        raise HTTPException(status_code=400, detail="Compare at most 6 products at a time")
    conn = get_conn()
    results = []
    for pid in ids:
        p = conn.execute("""
            SELECT p.*, m.name as manufacturer_name, m.category, m.confidence_tier,
                   m.origin, m.ksa_status, m.headquarters
            FROM products p JOIN manufacturers m ON m.id = p.manufacturer_id WHERE p.id = ?
        """, (pid,)).fetchone()
        if not p:
            continue
        p = dict(p)
        dists = conn.execute("""
            SELECT d.name FROM company_distributors cd
            JOIN distributors d ON d.id = cd.distributor_id
            WHERE cd.manufacturer_id = ?
        """, (p["manufacturer_id"],)).fetchall()
        p["distributors"] = [d["name"] for d in dists]
        try:
            p["specs"] = json.loads(p.get("specs_json") or "{}")
        except (ValueError, TypeError):
            p["specs"] = {}
        results.append(p)
    conn.close()
    if not results:
        raise HTTPException(status_code=404, detail="None of the requested product_ids were found")
    fields = ["product_name", "manufacturer_name", "origin", "category", "department", "product_type",
              "throughput", "sample_types", "certifications", "confidence_tier", "distributors", "description"]
    # Union of every structured spec key present across the compared products (in first-seen
    # order), so the comparison table gains a row for each real spec attribute (e.g. "Throughput
    # (single module)", "Cuvette capacity") without needing a fixed schema for every product type.
    spec_keys = []
    for p in results:
        for k in p["specs"].keys():
            if k not in spec_keys:
                spec_keys.append(k)
    return {"products": results, "comparison_fields": fields, "spec_fields": spec_keys, "count": len(results)}

@app.get("/companies/{company_id}/insight")
def company_insight(company_id: int):
    company = get_company(company_id)
    return {"company": company["name"], "insight": generate_insight(company), "generation_method": "rule-based template over verified fields — not a live AI-generated claim"}

@app.get("/regulatory")
def list_regulatory():
    """Regulatory status records. Deliberately sparse — only contains facts
    already verified with a source during product research. Companies/products
    not listed here simply have no regulatory claim recorded yet; that is not
    the same as 'not approved'."""
    conn = get_conn()
    rows = conn.execute("""
        SELECT r.*, m.name as manufacturer_name, p.product_name
        FROM regulatory_status r
        JOIN manufacturers m ON m.id = r.manufacturer_id
        LEFT JOIN products p ON p.id = r.product_id
        ORDER BY m.name
    """).fetchall()
    conn.close()
    return [dict(r) for r in rows]

@app.get("/companies/{company_id}/regulatory")
def company_regulatory(company_id: int):
    conn = get_conn()
    rows = conn.execute("""
        SELECT r.*, p.product_name FROM regulatory_status r
        LEFT JOIN products p ON p.id = r.product_id
        WHERE r.manufacturer_id = ?
    """, (company_id,)).fetchall()
    conn.close()
    if not rows:
        return {"company_id": company_id, "regulatory": [], "note": "No regulatory status recorded yet for this company — not yet researched, not a negative finding."}
    return {"company_id": company_id, "regulatory": [dict(r) for r in rows]}

@app.get("/knowledge-graph")
def knowledge_graph():
    """Graph view of verified relationships already in the database.
    Every node and edge here is derived from existing rows — nothing is
    invented for this endpoint. Company-distributor edges reuse the same
    LIKE-based matching rule as /companies/{id} for consistency."""
    conn = get_conn()
    companies = conn.execute("SELECT id, name, category, status_tag FROM manufacturers").fetchall()
    products = conn.execute("SELECT id, product_name, manufacturer_id FROM products").fetchall()
    techs = conn.execute("SELECT id, name FROM technologies").fetchall()
    mtech = conn.execute("SELECT manufacturer_id, technology_id FROM manufacturer_technology").fetchall()
    dists = conn.execute("SELECT id, name, represents FROM distributors").fetchall()
    conn.close()

    nodes = []
    edges = []
    for c in companies:
        nodes.append({"id": f"company-{c['id']}", "label": c["name"], "type": "company", "category": c["category"], "status": c["status_tag"]})
    for p in products:
        nodes.append({"id": f"product-{p['id']}", "label": p["product_name"], "type": "product"})
        edges.append({"source": f"company-{p['manufacturer_id']}", "target": f"product-{p['id']}", "relation": "makes"})
    for t in techs:
        nodes.append({"id": f"tech-{t['id']}", "label": t["name"], "type": "technology"})
    for mt in mtech:
        edges.append({"source": f"company-{mt['manufacturer_id']}", "target": f"tech-{mt['technology_id']}", "relation": "uses"})
    for d in dists:
        nodes.append({"id": f"dist-{d['id']}", "label": d["name"], "type": "distributor"})
        for c in companies:
            name_key = c["name"].split(" (")[0]
            if d["represents"] and name_key.lower() in d["represents"].lower():
                edges.append({"source": f"dist-{d['id']}", "target": f"company-{c['id']}", "relation": "distributes"})

    return {"nodes": nodes, "edges": edges, "node_count": len(nodes), "edge_count": len(edges),
            "note": "Every node/edge is derived from existing verified database rows — this is a graph view of real data, not a generated network."}

# ---------------- Patient management (Doctor AI) ----------------

class PatientCreate(BaseModel):
    full_name: str = Field(..., min_length=1)
    gender: str | None = None
    age: int | None = None
    department: str | None = None
    visit_type: str = "Outpatient"

class VisitCreate(BaseModel):
    visit_type: str | None = None
    notes: str | None = None

class PatientChatMessage(BaseModel):
    content: Any

def _gen_patient_code(conn) -> str:
    row = conn.execute("SELECT MAX(id) as m FROM patients").fetchone()
    next_id = (row["m"] or 10300) + 1
    return str(max(next_id, 10301))

@app.get("/patients")
def list_patients(patient_code: str | None = None, name: str | None = None, visit_type: str | None = None,
                   department: str | None = None, limit: int = 20, offset: int = 0):
    if limit < 1 or limit > 200:
        raise HTTPException(status_code=400, detail="limit must be between 1 and 200")
    conn = get_conn()
    query = "SELECT * FROM patients WHERE 1=1"
    params = []
    if patient_code:
        query += " AND patient_code LIKE ?"
        params.append(f"%{patient_code}%")
    if name:
        query += " AND full_name LIKE ?"
        params.append(f"%{name}%")
    if visit_type:
        query += " AND visit_type = ?"
        params.append(visit_type)
    if department:
        query += " AND department = ?"
        params.append(department)
    total = conn.execute(f"SELECT count(*) c FROM ({query})", params).fetchone()["c"]
    query += " ORDER BY id DESC LIMIT ? OFFSET ?"
    params += [limit, offset]
    rows = conn.execute(query, params).fetchall()

    results = []
    for r in rows:
        d = dict(r)
        last_visit = conn.execute(
            "SELECT visit_time FROM patient_visits WHERE patient_id = ? ORDER BY visit_time DESC LIMIT 1",
            (d["id"],)
        ).fetchone()
        d["last_visit_time"] = last_visit["visit_time"] if last_visit else d["created_at"]
        results.append(d)
    conn.close()
    return {"results": results, "total": total, "limit": limit, "offset": offset}

@app.post("/patients")
def create_patient(payload: PatientCreate):
    conn = get_conn()
    code = _gen_patient_code(conn)
    cur = conn.execute(
        "INSERT INTO patients (patient_code, full_name, gender, age, department, visit_type) VALUES (?,?,?,?,?,?)",
        (code, payload.full_name, payload.gender, payload.age, payload.department, payload.visit_type)
    )
    patient_id = cur.lastrowid
    conn.execute(
        "INSERT INTO patient_visits (patient_id, visit_type, notes) VALUES (?,?,?)",
        (patient_id, payload.visit_type, "Initial registration")
    )
    conn.commit()
    logger.info(f"POST /patients -> created patient {patient_id} ({payload.full_name})")
    conn.close()
    return {"id": patient_id, "patient_code": code, "status": "created"}

@app.get("/patients/{patient_id}")
def get_patient(patient_id: int):
    conn = get_conn()
    patient = conn.execute("SELECT * FROM patients WHERE id = ?", (patient_id,)).fetchone()
    if not patient:
        conn.close()
        raise HTTPException(status_code=404, detail="Patient not found")
    visits = conn.execute(
        "SELECT * FROM patient_visits WHERE patient_id = ? ORDER BY visit_time DESC", (patient_id,)
    ).fetchall()
    reports = conn.execute(
        "SELECT * FROM patient_reports WHERE patient_id = ? ORDER BY uploaded_at DESC", (patient_id,)
    ).fetchall()
    conn.close()
    return {
        "patient": dict(patient),
        "visits": [dict(v) for v in visits],
        "reports": [dict(r) for r in reports],
    }

@app.post("/patients/{patient_id}/visits")
def add_visit(patient_id: int, payload: VisitCreate):
    conn = get_conn()
    exists = conn.execute("SELECT id FROM patients WHERE id = ?", (patient_id,)).fetchone()
    if not exists:
        conn.close()
        raise HTTPException(status_code=404, detail="Patient not found")
    cur = conn.execute(
        "INSERT INTO patient_visits (patient_id, visit_type, notes) VALUES (?,?,?)",
        (patient_id, payload.visit_type, payload.notes)
    )
    conn.commit()
    visit_id = cur.lastrowid
    conn.close()
    return {"id": visit_id, "status": "created"}

@app.get("/patients/{patient_id}/chat")
def get_patient_chat(patient_id: int):
    conn = get_conn()
    exists = conn.execute("SELECT id FROM patients WHERE id = ?", (patient_id,)).fetchone()
    if not exists:
        conn.close()
        raise HTTPException(status_code=404, detail="Patient not found")
    rows = conn.execute(
        "SELECT role, content, created_at FROM patient_chat_messages WHERE patient_id = ? ORDER BY id ASC",
        (patient_id,)
    ).fetchall()
    conn.close()
    return {"messages": [dict(r) for r in rows]}

DOCTOR_AI_PATIENT_SYSTEM = (
    "You are Doctor AI, a clinical laboratory decision-support assistant on MedForsa GCC, used by physicians "
    "and laboratory professionals across Saudi Arabia and the Gulf. You are reviewing a specific patient's "
    "record and discussing it with the clinician. Patient context: {patient_context}\n\n"
    "Your role: (1) summarize and interpret uploaded lab values against standard reference ranges, "
    "(2) discuss differential considerations and relevant pathophysiology, (3) cite established literature, "
    "clinical guidelines, or textbooks by name where relevant, and (4) support an evidence-based discussion "
    "of the case using the patient's history across visits. You are strictly a reference and discussion tool, "
    "not a diagnostic authority: never state a definitive diagnosis as fact, always frame clinical observations "
    "as possibilities ('may suggest', 'is consistent with', 'differential includes') rather than certainties, "
    "and always make clear that final diagnosis and treatment decisions rest with the treating physician after "
    "direct clinical evaluation. For medication or dosing questions, give only general reference-level "
    "information and recommend consulting current prescribing references. Reply in the same language the "
    "user writes in (Arabic or English). Keep responses focused and clinically useful."
    "{lab_info_block}"
)

def _extract_text_for_matching(content):
    """Pull plain text out of a chat message's content, whether it's a plain string
    or a list of Anthropic content blocks (text/image/document)."""
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        parts = []
        for block in content:
            if isinstance(block, dict) and block.get("type") == "text":
                parts.append(block.get("text", ""))
        return " ".join(parts)
    return ""

def find_matching_lab_tests(conn, text, limit=5):
    """Look up MedForsa GCC's internal Lab Info reference library (lab_tests table)
    for any test named or aliased in the given text, so Doctor AI grounds its
    answer in our own sourced reference ranges instead of relying on memory alone."""
    text_lower = (text or "").lower()
    if not text_lower.strip():
        return []
    rows = conn.execute(
        "SELECT slug, name_en, name_ar, aliases, reference_ranges_json, reference_ranges_verified, "
        "clinical_significance_en, critical_values_en FROM lab_tests WHERE is_published = 1"
    ).fetchall()
    matches = []
    for r in rows:
        candidates = [r["name_en"], r["name_ar"]] + (r["aliases"].split(",") if r["aliases"] else [])
        for c in candidates:
            c = c.strip().lower()
            if c and len(c) >= 3 and c in text_lower:
                matches.append(r)
                break
        if len(matches) >= limit:
            break
    return matches

def build_lab_info_block(matches):
    if not matches:
        return ""
    lines = [
        "\n\nInternal Lab Info reference data (MedForsa GCC's own sourced test library -- "
        "prefer these values over general memory when discussing the tests below, and mention "
        "to the clinician that more detail is available on the platform's Lab Info page):"
    ]
    for m in matches:
        try:
            ranges = json.loads(m["reference_ranges_json"]) if m["reference_ranges_json"] else []
        except (ValueError, TypeError):
            ranges = []
        range_str = "; ".join(
            f"{rg.get('parameter','')} ({rg.get('population','')}): {rg.get('range','')}" for rg in ranges
        ) or "no structured ranges on file"
        verified = "verified" if m["reference_ranges_verified"] else "UNVERIFIED -- flag to clinician if used"
        crit_note = f" CRITICAL VALUE ALERT: {m['critical_values_en']}." if m["critical_values_en"] else ""
        lines.append(
            f"- {m['name_en']} / {m['name_ar']} [{verified}] -- {range_str}.{crit_note} "
            f"Full page: /lab-info?slug={m['slug']}"
        )
    return "\n".join(lines)

@app.post("/patients/{patient_id}/chat")
def send_patient_chat(patient_id: int, payload: PatientChatMessage):
    conn = get_conn()
    patient = conn.execute("SELECT * FROM patients WHERE id = ?", (patient_id,)).fetchone()
    if not patient:
        conn.close()
        raise HTTPException(status_code=404, detail="Patient not found")

    history_rows = conn.execute(
        "SELECT role, content FROM patient_chat_messages WHERE patient_id = ? ORDER BY id ASC",
        (patient_id,)
    ).fetchall()

    import json as _json
    messages = []
    for r in history_rows:
        content = r["content"]
        try:
            content = _json.loads(content)
        except (ValueError, TypeError):
            pass
        messages.append({"role": r["role"], "content": content})
    messages.append({"role": "user", "content": payload.content})

    patient_context = (
        f"{patient['full_name']}, {patient['age']}-year-old {patient['gender'] or 'unknown gender'}, "
        f"department: {patient['department'] or 'not specified'}, visit type: {patient['visit_type']}."
    )

    latest_text = _extract_text_for_matching(payload.content)
    lab_matches = find_matching_lab_tests(conn, latest_text)
    lab_info_block = build_lab_info_block(lab_matches)

    system_prompt = DOCTOR_AI_PATIENT_SYSTEM.format(
        patient_context=patient_context, lab_info_block=lab_info_block
    )

    result = call_anthropic(system_prompt, messages)
    reply_text = "\n".join(b.get("text", "") for b in result.get("content", []) if b.get("type") == "text")

    user_content_store = payload.content if isinstance(payload.content, str) else _json.dumps(payload.content)
    conn.execute(
        "INSERT INTO patient_chat_messages (patient_id, role, content) VALUES (?,?,?)",
        (patient_id, "user", user_content_store)
    )
    conn.execute(
        "INSERT INTO patient_chat_messages (patient_id, role, content) VALUES (?,?,?)",
        (patient_id, "assistant", reply_text)
    )
    conn.commit()
    conn.close()
    logger.info(f"POST /patients/{patient_id}/chat -> exchanged message")
    return {"reply": reply_text}

# ---------------- Lab Test Reference ("دليل التحاليل المخبرية") ----------------

class LabTestCreate(BaseModel):
    slug: str = Field(..., min_length=1, max_length=100)
    name_en: str = Field(..., min_length=1)
    name_ar: str | None = None  # Lab Info content is English-only going forward; kept for legacy compatibility
    aliases: str | None = None
    category: str = Field(..., min_length=1)
    purpose_en: str | None = None
    purpose_ar: str | None = None
    specimen_type: str | None = None
    collection_notes_en: str | None = None
    collection_notes_ar: str | None = None
    methodology_en: str | None = None
    methodology_ar: str | None = None
    reference_ranges: list[dict] | None = None
    reference_ranges_verified: bool = False
    clinical_significance_en: str | None = None
    clinical_significance_ar: str | None = None
    critical_values_en: str | None = None
    interfering_factors_en: str | None = None
    questions_to_ask_en: str | None = None
    next_steps_en: str | None = None
    associated_conditions: list[dict] | None = None
    related_tests: list[str] | None = None
    sources: list[dict] = Field(..., min_length=1)
    is_published: bool = True

def _lab_test_row_to_dict(row, include_sources=False):
    d = dict(row)
    for json_field, key in [("reference_ranges_json", "reference_ranges"),
                             ("associated_conditions_json", "associated_conditions"),
                             ("sources_json", "sources")]:
        raw = d.pop(json_field, None)
        try:
            d[key] = json.loads(raw) if raw else []
        except (ValueError, TypeError):
            d[key] = []
    d.pop("related_tests_json", None)
    if not include_sources:
        # Sources are kept for MedForsa GCC's own internal verification/audit trail
        # but are not surfaced in the public-facing Lab Info UI.
        d.pop("sources", None)
    d["reference_ranges_verified"] = bool(d.get("reference_ranges_verified"))
    d["is_published"] = bool(d.get("is_published"))
    return d

def _resolve_related_tests(conn, row, limit=5):
    """Manually curated related_tests_json (list of slugs) takes priority; falls back to
    other published tests in the same category when no curation exists yet."""
    try:
        related_slugs = json.loads(row["related_tests_json"]) if row["related_tests_json"] else []
    except (ValueError, TypeError):
        related_slugs = []

    results = []
    seen = {row["slug"]}
    if related_slugs:
        placeholders = ",".join("?" for _ in related_slugs)
        rows = conn.execute(
            f"SELECT slug, name_en, category FROM lab_tests WHERE slug IN ({placeholders}) AND is_published = 1",
            related_slugs
        ).fetchall()
        by_slug = {r["slug"]: r for r in rows}
        for s in related_slugs:
            if s in by_slug and s not in seen:
                r = by_slug[s]
                results.append({"slug": r["slug"], "name_en": r["name_en"], "category": r["category"]})
                seen.add(s)

    if len(results) < limit:
        fallback_rows = conn.execute(
            "SELECT slug, name_en, category FROM lab_tests WHERE category = ? AND is_published = 1 AND slug != ? ORDER BY name_en LIMIT ?",
            (row["category"], row["slug"], limit - len(results))
        ).fetchall()
        for r in fallback_rows:
            if r["slug"] not in seen:
                results.append({"slug": r["slug"], "name_en": r["name_en"], "category": r["category"]})
                seen.add(r["slug"])

    return results[:limit]

@app.get("/lab-tests-list")
def list_lab_tests(q: str | None = None, category: str | None = None,
                    limit: int = 50, offset: int = 0, include_unpublished: bool = False):
    if limit < 1 or limit > 200:
        raise HTTPException(status_code=400, detail="limit must be between 1 and 200")
    conn = get_conn()
    query = "SELECT * FROM lab_tests WHERE 1=1"
    params = []
    if not include_unpublished:
        query += " AND is_published = 1"
    if category:
        query += " AND category = ?"
        params.append(category)
    if q:
        query += " AND (name_en LIKE ? OR name_ar LIKE ? OR aliases LIKE ?)"
        like = f"%{q}%"
        params += [like, like, like]
    total = conn.execute(f"SELECT count(*) c FROM ({query})", params).fetchone()["c"]
    query += " ORDER BY name_en LIMIT ? OFFSET ?"
    params += [limit, offset]
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return {"total": total, "limit": limit, "offset": offset,
            "results": [_lab_test_row_to_dict(r) for r in rows]}

@app.get("/lab-tests-categories")
def lab_test_categories():
    conn = get_conn()
    rows = conn.execute(
        "SELECT category, count(*) c FROM lab_tests WHERE is_published = 1 GROUP BY category ORDER BY category"
    ).fetchall()
    conn.close()
    return {"categories": [dict(r) for r in rows]}

@app.get("/lab-tests-detail/{slug}")
def get_lab_test(slug: str):
    conn = get_conn()
    row = conn.execute("SELECT * FROM lab_tests WHERE slug = ?", (slug,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Lab test not found")
    result = _lab_test_row_to_dict(row, include_sources=False)
    result["related_tests"] = _resolve_related_tests(conn, row)
    conn.close()
    return result

@app.post("/lab-tests-list")
def create_lab_test(payload: LabTestCreate):
    conn = get_conn()
    existing = conn.execute("SELECT id FROM lab_tests WHERE slug = ?", (payload.slug,)).fetchone()
    if existing:
        conn.close()
        raise HTTPException(status_code=409, detail="A lab test with this slug already exists")
    cur = conn.execute(
        """INSERT INTO lab_tests
        (slug, name_en, name_ar, aliases, category, purpose_en, purpose_ar, specimen_type,
         collection_notes_en, collection_notes_ar, methodology_en, methodology_ar,
         reference_ranges_json, reference_ranges_verified, clinical_significance_en, clinical_significance_ar,
         critical_values_en, interfering_factors_en, questions_to_ask_en, next_steps_en,
         associated_conditions_json, related_tests_json, sources_json, is_published)
        VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""",
        (payload.slug, payload.name_en, payload.name_ar or "", payload.aliases, payload.category,
         payload.purpose_en, payload.purpose_ar, payload.specimen_type,
         payload.collection_notes_en, payload.collection_notes_ar,
         payload.methodology_en, payload.methodology_ar,
         json.dumps(payload.reference_ranges or []), int(payload.reference_ranges_verified),
         payload.clinical_significance_en, payload.clinical_significance_ar,
         payload.critical_values_en, payload.interfering_factors_en,
         payload.questions_to_ask_en, payload.next_steps_en,
         json.dumps(payload.associated_conditions or []), json.dumps(payload.related_tests or []),
         json.dumps(payload.sources), int(payload.is_published))
    )
    conn.commit()
    new_id = cur.lastrowid
    log_action(conn, "lab_tests", new_id, "create", f"Added lab test: {payload.name_en}")
    conn.close()
    logger.info(f"POST /lab-tests-list -> created lab test '{payload.slug}'")
    return {"id": new_id, "slug": payload.slug, "status": "created"}

# ---------------- Order Sets ("مجموعات التحاليل") ----------------

class OrderSetCreate(BaseModel):
    slug: str = Field(..., min_length=1, max_length=100)
    name_en: str = Field(..., min_length=1)
    description_en: str | None = None
    category: str | None = None
    test_slugs: list[str] = Field(..., min_length=1)
    is_published: bool = True

@app.get("/order-sets-list")
def list_order_sets():
    conn = get_conn()
    rows = conn.execute(
        "SELECT * FROM order_sets WHERE is_published = 1 ORDER BY name_en"
    ).fetchall()
    results = []
    for row in rows:
        d = dict(row)
        try:
            slugs = json.loads(d.pop("test_slugs_json") or "[]")
        except (ValueError, TypeError):
            slugs = []
        d["is_published"] = bool(d.get("is_published"))
        # Resolve each member test to its name for a quick preview on the card,
        # without needing the client to make N follow-up detail calls.
        tests = []
        if slugs:
            placeholders = ",".join("?" for _ in slugs)
            test_rows = conn.execute(
                f"SELECT slug, name_en, category FROM lab_tests WHERE slug IN ({placeholders}) AND is_published = 1",
                slugs
            ).fetchall()
            by_slug = {r["slug"]: r for r in test_rows}
            for s in slugs:
                if s in by_slug:
                    r = by_slug[s]
                    tests.append({"slug": r["slug"], "name_en": r["name_en"], "category": r["category"]})
        d["tests"] = tests
        d["test_count"] = len(tests)
        results.append(d)
    conn.close()
    return {"results": results}

@app.get("/order-sets-detail/{slug}")
def get_order_set(slug: str):
    conn = get_conn()
    row = conn.execute("SELECT * FROM order_sets WHERE slug = ?", (slug,)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Order set not found")
    d = dict(row)
    try:
        slugs = json.loads(d.pop("test_slugs_json") or "[]")
    except (ValueError, TypeError):
        slugs = []
    d["is_published"] = bool(d.get("is_published"))
    tests = []
    if slugs:
        placeholders = ",".join("?" for _ in slugs)
        test_rows = conn.execute(
            f"SELECT slug, name_en, category, purpose_en, specimen_type FROM lab_tests WHERE slug IN ({placeholders}) AND is_published = 1",
            slugs
        ).fetchall()
        by_slug = {r["slug"]: dict(r) for r in test_rows}
        for s in slugs:
            if s in by_slug:
                tests.append(by_slug[s])
    d["tests"] = tests
    conn.close()
    return d

@app.post("/order-sets-list")
def create_order_set(payload: OrderSetCreate):
    conn = get_conn()
    existing = conn.execute("SELECT id FROM order_sets WHERE slug = ?", (payload.slug,)).fetchone()
    if existing:
        conn.close()
        raise HTTPException(status_code=409, detail="An order set with this slug already exists")
    cur = conn.execute(
        "INSERT INTO order_sets (slug, name_en, description_en, category, test_slugs_json, is_published) VALUES (?,?,?,?,?,?)",
        (payload.slug, payload.name_en, payload.description_en, payload.category,
         json.dumps(payload.test_slugs), int(payload.is_published))
    )
    conn.commit()
    new_id = cur.lastrowid
    conn.close()
    logger.info(f"POST /order-sets-list -> created order set '{payload.slug}'")
    return {"id": new_id, "slug": payload.slug, "status": "created"}

# ---------------- WhatsApp Assistant (Twilio) ----------------

WHATSAPP_SYSTEM_PROMPT = (
    "You are the MedForsa GCC WhatsApp assistant, helping physicians, laboratory professionals, "
    "and IVD distribution contacts across Saudi Arabia and the Gulf with quick questions about "
    "lab tests, reference ranges, and clinical significance. You draw on MedForsa GCC's own Lab "
    "Info reference library, which is provided to you below when relevant -- prefer that sourced "
    "data over general memory whenever it's given to you.\n\n"
    "Keep replies concise and WhatsApp-appropriate: short paragraphs, no long tables, no markdown "
    "headers. You are a reference and discussion tool, not a diagnostic authority -- never state a "
    "definitive diagnosis, frame clinical points as possibilities ('may suggest', 'is consistent "
    "with') not certainties, and make clear that direct clinical evaluation and the treating "
    "physician's judgment take priority. For medication/dosing questions, give only general "
    "reference-level information. Reply in the same language the user writes in (Arabic or English)."
    "{lab_info_block}"
)

def _twiml_response(message_text: str) -> Response:
    import xml.sax.saxutils as saxutils
    escaped = saxutils.escape(message_text or "")
    xml = f'<?xml version="1.0" encoding="UTF-8"?><Response><Message>{escaped}</Message></Response>'
    return Response(content=xml, media_type="application/xml")

def _validate_twilio_signature(request: Request, form_params: dict) -> bool:
    """Verifies the X-Twilio-Signature header per Twilio's documented HMAC-SHA1
    scheme, so this endpoint can't be spoofed/abused by arbitrary POSTs once
    it's live (each request otherwise costs a real Anthropic API call).
    Returns True (allow) if TWILIO_AUTH_TOKEN isn't set yet -- same ready-but-
    needs-credentials pattern as the rest of the Twilio integration; once the
    token is configured this becomes a hard requirement."""
    auth_token = os.environ.get("TWILIO_AUTH_TOKEN")
    if not auth_token:
        return True
    signature = request.headers.get("X-Twilio-Signature", "")
    url = str(request.url)
    data = url + "".join(f"{k}{v}" for k, v in sorted(form_params.items()))
    expected = base64.b64encode(
        hmac.new(auth_token.encode(), data.encode(), hashlib.sha1).digest()
    ).decode()
    return hmac.compare_digest(expected, signature)

@app.post("/whatsapp-webhook")
async def whatsapp_webhook(request: Request):
    """Twilio WhatsApp inbound webhook. Twilio POSTs form-encoded fields
    (Body, From, WaId, ...) for every inbound message; we reply with TwiML."""
    form = await request.form()
    if not _validate_twilio_signature(request, dict(form)):
        logger.warning("WhatsApp webhook: rejected request with invalid Twilio signature")
        raise HTTPException(status_code=403, detail="Invalid signature")
    body_text = (form.get("Body") or "").strip()
    from_number = form.get("From") or form.get("WaId") or "unknown"

    if not body_text:
        return _twiml_response("I didn't receive any text in that message -- please send your question as text.")

    conn = get_conn()
    history_rows = conn.execute(
        "SELECT role, content FROM whatsapp_messages WHERE phone_number = ? ORDER BY id DESC LIMIT 12",
        (from_number,)
    ).fetchall()
    history_rows = list(reversed(history_rows))
    messages = [{"role": r["role"], "content": r["content"]} for r in history_rows]
    messages.append({"role": "user", "content": body_text})

    lab_matches = find_matching_lab_tests(conn, body_text)
    lab_info_block = build_lab_info_block(lab_matches)
    system_prompt = WHATSAPP_SYSTEM_PROMPT.format(lab_info_block=lab_info_block)

    try:
        result = call_anthropic(system_prompt, messages)
        reply_text = "".join(
            block.get("text", "") for block in result.get("content", []) if block.get("type") == "text"
        ).strip() or "Sorry, I couldn't generate a response just now. Please try again shortly."
    except Exception as e:
        logger.error(f"WhatsApp webhook call_anthropic failed: {e}")
        reply_text = "Sorry, I'm having trouble reaching the assistant right now. Please try again shortly."

    conn.execute(
        "INSERT INTO whatsapp_messages (phone_number, role, content) VALUES (?,?,?)",
        (from_number, "user", body_text)
    )
    conn.execute(
        "INSERT INTO whatsapp_messages (phone_number, role, content) VALUES (?,?,?)",
        (from_number, "assistant", reply_text)
    )
    conn.commit()
    conn.close()
    logger.info(f"POST /whatsapp-webhook -> replied to {from_number}")
    return _twiml_response(reply_text)

# ---------------- Price List Parser (for the Deal & TCO Calculator) ----------------

PRICE_KEYWORDS = ["price", "cost", "unit price", "سعر", "التكلفه", "التكلفة", "قيمه", "قيمة"]
QTY_KEYWORDS = ["test", "tests", "qty", "quantity", "pack", "kit size", "size", "اختبار",
                "اختبارات", "عدد", "عبوه", "عبوة", "حجم"]
NAME_KEYWORDS = ["name", "product", "item", "kit", "description", "اسم", "منتج", "الصنف", "الكيت"]
# Columns matching these should never be guessed as price or tests-per-kit --
# catalog/reference numbers look numeric but aren't a usable quantity or price.
EXCLUDE_KEYWORDS = ["code", "catalog", "cat.", "cat no", "sku", "ref", "reference", "item no",
                    "part no", "id", "s.no", "sl.no", "sl no", "s/n", "sr.no", "sr no",
                    "serial", "seq", "no.", "كود", "رقم الصنف", "الرقم المرجعي", "م."]

def _guess_column(headers, keywords):
    """Best-effort column detection: returns the header string whose lowercase
    text contains one of the given keywords, or None if no match is found."""
    for h in headers:
        h_norm = str(h).strip().lower()
        for kw in keywords:
            if kw.lower() in h_norm:
                return h
    return None

def _extract_numeric(value):
    """Pulls a float out of a cell that might contain currency symbols, commas,
    or Arabic digits mixed with text (e.g. 'SAR 1,200.00' -> 1200.0)."""
    import re
    if value is None:
        return None
    s = str(value).strip()
    if not s:
        return None
    s = s.replace(",", "").replace("SAR", "").replace("ريال", "").strip()
    m = re.search(r"-?\d+(?:\.\d+)?", s)
    return float(m.group()) if m else None

def _rows_from_dataframe(df):
    """Given a pandas DataFrame with a header row, best-effort extract
    {name, price, tests_per_kit} for each row using keyword-matched columns,
    falling back to the first two numeric columns if no keyword match is found.
    Excludes catalog/product-code columns (numeric-looking but not a usable
    price or quantity) and sanity-checks the tests-per-kit column so an
    implausible value (e.g. a 9-digit product code) is left blank rather than
    shown as if it were real data -- the frontend then leaves that cell empty
    for the user to fill in manually instead of misleading them."""
    import pandas as pd
    headers = list(df.columns)
    is_excluded = lambda h: any(kw in str(h).strip().lower() for kw in EXCLUDE_KEYWORDS)
    usable_headers = [h for h in headers if not is_excluded(h)]

    name_col = _guess_column(usable_headers, NAME_KEYWORDS)
    price_col = _guess_column(usable_headers, PRICE_KEYWORDS)
    qty_col = _guess_column(usable_headers, QTY_KEYWORDS)

    numeric_cols = []
    if price_col is None or qty_col is None:
        for h in usable_headers:
            try:
                if pd.to_numeric(df[h], errors="coerce").notna().sum() > 0:
                    numeric_cols.append(h)
            except Exception:
                continue

    if price_col is None and numeric_cols:
        price_col = numeric_cols[0]
    if qty_col is None:
        remaining = [c for c in numeric_cols if c != price_col]
        if remaining:
            qty_col = remaining[0]

    # Sanity check 1: a real "tests per kit" value is realistically 1-5,000.
    # If the candidate qty column's values are mostly outside that range (e.g.
    # a product-code column that slipped through), don't use it -- leave
    # tests_per_kit blank rather than show a misleading number.
    if qty_col is not None:
        numeric_qty = pd.to_numeric(df[qty_col], errors="coerce").dropna()
        if len(numeric_qty) > 0:
            plausible_fraction = ((numeric_qty >= 1) & (numeric_qty <= 5000)).mean()
            if plausible_fraction < 0.5:
                qty_col = None

    # Sanity check 2: a row-index / serial-number column (1, 2, 3, 4, ...) can
    # slip through as "numeric with no keyword match" regardless of its header
    # text -- detect this by checking if the values are just a consecutive
    # integer run starting at 0 or 1, and reject it if so.
    if qty_col is not None:
        numeric_qty = pd.to_numeric(df[qty_col], errors="coerce").dropna().tolist()
        if len(numeric_qty) >= 3:
            sorted_vals = sorted(numeric_qty)
            is_sequential = all(
                float(v) == float(sorted_vals[0]) + i for i, v in enumerate(sorted_vals)
            )
            starts_at_index = sorted_vals[0] in (0, 1)
            if is_sequential and starts_at_index:
                qty_col = None

    if name_col is None:
        non_numeric = [h for h in usable_headers if h not in (price_col, qty_col)]
        name_col = non_numeric[0] if non_numeric else (usable_headers[0] if usable_headers else headers[0])

    results = []
    for _, row in df.iterrows():
        name = str(row.get(name_col, "")).strip() if name_col else ""
        price = _extract_numeric(row.get(price_col)) if price_col else None
        qty = _extract_numeric(row.get(qty_col)) if qty_col else None
        if not name or name.lower() == "nan":
            continue
        if price is None and qty is None:
            continue
        results.append({"name": name, "price": price, "tests_per_kit": qty})
    return results

@app.post("/parse-price-list")
async def parse_price_list(file: UploadFile = File(...)):
    """Accepts a PDF or Excel price list and returns a best-effort extraction
    of {name, price, tests_per_kit} line items, for the Deal & TCO Calculator's
    'Upload Price List' feature. Extraction is heuristic (keyword-matched
    columns with a numeric-column fallback) since real-world price list layouts
    vary widely -- the frontend lets the user review and pick the right row
    rather than trusting the parse blindly."""
    filename = (file.filename or "").lower()
    content = await file.read()
    if len(content) > 15 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 15MB)")

    import io
    import pandas as pd
    items = []
    warnings = []

    try:
        if filename.endswith((".xlsx", ".xls")):
            xls = pd.ExcelFile(io.BytesIO(content))
            for sheet_name in xls.sheet_names:
                df = xls.parse(sheet_name)
                df = df.dropna(axis=0, how="all")
                if df.empty:
                    continue
                sheet_items = _rows_from_dataframe(df)
                items.extend(sheet_items)
            if not items:
                warnings.append("No recognizable price/quantity columns found in the spreadsheet.")

        elif filename.endswith(".pdf"):
            import pdfplumber
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                found_table = False
                for page in pdf.pages:
                    tables = page.extract_tables()
                    for table in tables:
                        if not table or len(table) < 2:
                            continue
                        header, rows = table[0], table[1:]
                        df = pd.DataFrame(rows, columns=header)
                        items.extend(_rows_from_dataframe(df))
                        found_table = True
                if not found_table:
                    warnings.append("No tables detected in the PDF -- try exporting the price list as Excel/CSV for more reliable extraction, or enter the kit price and tests-per-kit manually.")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type -- please upload a .xlsx, .xls, or .pdf file")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"parse_price_list failed for {filename}: {e}")
        raise HTTPException(status_code=422, detail=f"Could not parse this file: {e}")

    logger.info(f"POST /parse-price-list -> extracted {len(items)} line items from {filename}")
    return {"filename": file.filename, "items": items[:200], "warnings": warnings}

# ---------------- Quotation Generator (brochure parsing + letterhead + docx output) ----------------

@app.post("/parse-brochure")
async def parse_brochure(file: UploadFile = File(...)):
    """Extracts a device name, manufacturer, an overview/feature bullet list,
    and a specs table from an uploaded product brochure PDF, for pre-filling
    the Quotation Generator. Heuristic and PDF-only (brochures are almost
    always distributed as PDF) -- the frontend lets the user review/edit
    everything before generating the final quotation."""
    filename = (file.filename or "").lower()
    if not filename.endswith(".pdf"):
        raise HTTPException(status_code=400, detail="Please upload a PDF brochure")
    content = await file.read()
    if len(content) > 20 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 20MB)")

    import io
    import pandas as pd
    import pdfplumber

    device_name = ""
    manufacturer = ""
    overview_bullets = []
    specs = []

    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            first_page_text = pdf.pages[0].extract_text() or "" if pdf.pages else ""
            lines = [l.strip() for l in first_page_text.split("\n") if l.strip()]
            # Heuristic: the device name is usually one of the first few short,
            # title-cased lines on page 1 that isn't a company name/address/date.
            skip_markers = ["attieh", "jeddah", "saudi", "rfq", "ref", "page", "attn", "2026", "2025", "2027"]
            for l in lines[:12]:
                if len(l) < 60 and not any(m in l.lower() for m in skip_markers) and any(c.isalpha() for c in l):
                    device_name = l
                    break

            # Manufacturer can appear on any page (often near the pricing table) --
            # search every page's text for a "Manufacturer:" line or a known brand name.
            known_brands = ["snibe", "roche", "abbott", "sysmex", "mindray", "siemens", "beckman",
                             "bio-rad", "biorad", "ortho", "bd ", "diasorin", "werfen"]
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                for l in page_text.split("\n"):
                    l = l.strip()
                    if "manufacturer" in l.lower():
                        manufacturer = l.split(":", 1)[-1].strip() if ":" in l else l
                        break
                    if any(l.lower().startswith(b) for b in known_brands):
                        manufacturer = l
                        break
                if manufacturer:
                    break

            noise_markers = ["attieh", "jeddah", "saudi arabia", "rfq", "our ref", "no of page",
                              "attn:", "www.", "email", "unified number", "paid-up capital",
                              device_name.lower()] if device_name else \
                             ["attieh", "jeddah", "saudi arabia", "rfq", "our ref", "no of page",
                              "attn:", "www.", "email", "unified number", "paid-up capital"]
            import re as _re
            date_pattern = _re.compile(r"\b\d{1,2}\s*(st|nd|rd|th)?\s*(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)", _re.I)

            for page in pdf.pages[:2]:
                text = page.extract_text() or ""
                for l in text.split("\n"):
                    l = l.strip()
                    if len(l) <= 15 or len(l) >= 140:
                        continue
                    if l.lower().startswith(("category", "attieh", "www.", "email")):
                        continue
                    if any(m in l.lower() for m in noise_markers):
                        continue
                    if date_pattern.search(l):
                        continue
                    if l in overview_bullets or any(k in l for k in ["|", "SAR", "C/N"]):
                        continue
                    overview_bullets.append(l)
                if len(overview_bullets) > 20:
                    break
            overview_bullets = overview_bullets[:14]

            for page in pdf.pages:
                for table in page.extract_tables():
                    if not table or len(table) < 2:
                        continue
                    header = [str(h or "").strip().lower() for h in table[0]]
                    if any("categor" in h for h in header) or any("detail" in h for h in header):
                        for row in table[1:]:
                            if len(row) >= 2 and row[0] and row[1]:
                                specs.append({"key": str(row[0]).strip(), "value": str(row[1]).strip()})
    except Exception as e:
        logger.error(f"parse_brochure failed for {filename}: {e}")
        raise HTTPException(status_code=422, detail=f"Could not parse this brochure: {e}")

    logger.info(f"POST /parse-brochure -> extracted device='{device_name}', {len(specs)} specs from {filename}")
    return {
        "filename": file.filename,
        "device_name": device_name,
        "manufacturer": manufacturer,
        "overview": overview_bullets,
        "specs": specs,
    }

@app.post("/upload-letterhead")
async def upload_letterhead(file: UploadFile = File(...)):
    """Stores the company logo (as base64 in the database, so it survives
    redeploys the same way all other MedForsa GCC data does) for reuse in
    generated quotations."""
    content = await file.read()
    if len(content) > 3 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="Logo file too large (max 3MB)")
    ext = (file.filename or "").split(".")[-1].lower()
    if ext not in ("png", "jpg", "jpeg"):
        raise HTTPException(status_code=400, detail="Please upload a PNG or JPG logo")
    b64 = base64.b64encode(content).decode("ascii")
    conn = get_conn()
    conn.execute(
        "INSERT INTO company_settings (key, value) VALUES ('letterhead_logo', ?) "
        "ON CONFLICT(key) DO UPDATE SET value = excluded.value",
        (f"data:image/{ext};base64,{b64}",)
    )
    conn.execute(
        "INSERT INTO company_settings (key, value) VALUES ('letterhead_logo_ext', ?) "
        "ON CONFLICT(key) DO UPDATE SET value = excluded.value",
        (ext,)
    )
    conn.commit()
    conn.close()
    logger.info("POST /upload-letterhead -> logo saved")
    return {"status": "saved"}

@app.get("/company-settings")
def get_company_settings():
    conn = get_conn()
    rows = conn.execute("SELECT key, value FROM company_settings").fetchall()
    conn.close()
    settings = {r["key"]: r["value"] for r in rows}
    return {"has_logo": "letterhead_logo" in settings, "company_name": settings.get("company_name", "")}

class QuotationRequest(BaseModel):
    device_name: str
    manufacturer: str | None = None
    overview: list[str] = []
    specs: list[dict] = []
    client_name: str | None = None
    client_location: str | None = None
    quotation_ref: str | None = None
    years: int = 3
    total_tests: int = 0
    price_per_test: float = 0.0
    total_deal_value: float = 0.0
    vat_pct: float = 15.0
    validity: str = "2 months from the date of quotation."
    delivery: str = "30 days from the date of firm Purchase Order"
    warranty: str = "As per manufacturer's standard warranty terms."
    installation: str = "By certified service engineer on site."
    power_rating: str = "220V/60Hz"
    payment: str = "Payment will be arranged under the reagent deal."
    prepared_by: str | None = None
    prepared_by_title: str | None = None
    approved_by: str | None = None
    approved_by_title: str | None = None

@app.post("/generate-quotation")
def generate_quotation(payload: QuotationRequest):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import io

    conn = get_conn()
    logo_row = conn.execute("SELECT value FROM company_settings WHERE key = 'letterhead_logo'").fetchone()
    conn.close()

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    navy = RGBColor(0x1A, 0x2B, 0x6B)

    if logo_row and logo_row["value"]:
        try:
            header, b64data = logo_row["value"].split(",", 1)
            logo_bytes = base64.b64decode(b64data)
            p = doc.add_paragraph()
            run = p.add_run()
            run.add_picture(io.BytesIO(logo_bytes), width=Inches(1.6))
        except Exception as e:
            logger.warning(f"Could not embed letterhead logo: {e}")

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_p.add_run(datetime.utcnow().strftime("%d %b %Y"))
    date_run.font.size = Pt(10); date_run.font.bold = True

    if payload.quotation_ref:
        ref_p = doc.add_paragraph()
        ref_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        ref_run = ref_p.add_run(f"Our Ref.: {payload.quotation_ref}")
        ref_run.font.size = Pt(9.5)

    if payload.client_name:
        client_p = doc.add_paragraph()
        client_run = client_p.add_run(payload.client_name)
        client_run.font.bold = True; client_run.font.size = Pt(11)
        if payload.client_location:
            loc_p = doc.add_paragraph()
            loc_run = loc_p.add_run(payload.client_location)
            loc_run.font.size = Pt(10)

    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(payload.device_name)
    title_run.font.bold = True; title_run.font.size = Pt(15); title_run.font.color.rgb = navy
    title_run.font.underline = True
    if payload.manufacturer:
        mfr_p = doc.add_paragraph()
        mfr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mfr_run = mfr_p.add_run(payload.manufacturer)
        mfr_run.font.size = Pt(10.5); mfr_run.font.italic = True

    if payload.overview:
        doc.add_paragraph()
        ov_head = doc.add_paragraph()
        ov_head_run = ov_head.add_run("Overview")
        ov_head_run.font.bold = True; ov_head_run.font.size = Pt(12)
        for b in payload.overview:
            bullet_p = doc.add_paragraph(style="List Bullet")
            bullet_run = bullet_p.add_run(b)
            bullet_run.font.size = Pt(10)

    if payload.specs:
        doc.add_paragraph()
        spec_head = doc.add_paragraph()
        spec_head_run = spec_head.add_run("Summary")
        spec_head_run.font.bold = True; spec_head_run.font.size = Pt(12)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = table.rows[0].cells
        hdr[0].text = "Category"; hdr[1].text = "Details"
        for cell in hdr:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.bold = True; r.font.size = Pt(9.5)
        for spec in payload.specs:
            row = table.add_row().cells
            row[0].text = str(spec.get("key", ""))
            row[1].text = str(spec.get("value", ""))
            for cell in row:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(9.5)

    doc.add_paragraph()
    deal_head = doc.add_paragraph()
    deal_head_run = deal_head.add_run(f"Contract will be ready after acceptance the prices below — {payload.device_name} Reagent Deal")
    deal_head_run.font.bold = True; deal_head_run.font.size = Pt(11)
    price_note = doc.add_paragraph()
    price_note_run = price_note.add_run(f"Price / test: {payload.price_per_test:.2f} SAR")
    price_note_run.font.bold = True; price_note_run.font.size = Pt(10.5)

    ptable = doc.add_table(rows=1, cols=5)
    ptable.style = "Table Grid"
    phdr = ptable.rows[0].cells
    for i, label in enumerate(["C/N", "Description", "Qty", "U/Price", "T/Price (SAR)"]):
        phdr[i].text = label
        for p in phdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(9.5)

    r1 = ptable.add_row().cells
    r1[0].text = "1"; r1[1].text = f"{payload.device_name} — Fully automated analyzer with computer and printer"
    r1[2].text = "1"; r1[3].text = "FOC"; r1[4].text = "FOC"

    r2 = ptable.add_row().cells
    r2[0].text = "2"
    r2[1].text = f"{payload.total_tests:,} tests / {payload.years} years as a reagent deal"
    r2[2].text = f"{payload.total_tests:,}"
    r2[3].text = f"{payload.price_per_test:.2f}"
    r2[4].text = f"{payload.total_deal_value:,.2f}"

    vat_amount = payload.total_deal_value * (payload.vat_pct / 100)
    r3 = ptable.add_row().cells
    r3[0].text = ""; r3[1].text = f"VAT {payload.vat_pct:.0f}%"
    r3[2].text = ""; r3[3].text = ""; r3[4].text = f"{vat_amount:,.2f}"

    r4 = ptable.add_row().cells
    r4[0].text = ""; r4[1].text = "Net Total (SAR)"
    r4[2].text = ""; r4[3].text = ""
    net_total_run_cell = r4[4]
    net_total_run_cell.text = f"{(payload.total_deal_value + vat_amount):,.2f}"
    for p in r4[1].paragraphs + net_total_run_cell.paragraphs:
        for r in p.runs:
            r.font.bold = True

    doc.add_paragraph()
    terms = [
        ("Validity", payload.validity),
        ("Delivery", payload.delivery),
        ("Warranty", payload.warranty),
        ("Installation", payload.installation),
        ("Power Rating", payload.power_rating),
        ("Payment", payload.payment),
    ]
    ttable = doc.add_table(rows=0, cols=2)
    ttable.style = "Table Grid"
    for label, value in terms:
        row = ttable.add_row().cells
        row[0].text = label; row[1].text = value
        for p in row[0].paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(9.5)
        for p in row[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9.5)

    doc.add_paragraph()
    sig_table = doc.add_table(rows=1, cols=2)
    sig_left, sig_right = sig_table.rows[0].cells
    if payload.prepared_by:
        p = sig_left.paragraphs[0]
        p.add_run("Prepared by:\n").font.bold = True
        p.add_run(payload.prepared_by + "\n")
        if payload.prepared_by_title:
            p.add_run(payload.prepared_by_title)
    if payload.approved_by:
        p = sig_right.paragraphs[0]
        p.add_run("Approved by:\n").font.bold = True
        p.add_run(payload.approved_by + "\n")
        if payload.approved_by_title:
            p.add_run(payload.approved_by_title)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    safe_name = "".join(c for c in payload.device_name if c.isalnum() or c in " -_").strip()
    safe_name = safe_name.replace(" ", "_")[:60] or "quotation"
    logger.info(f"POST /generate-quotation -> generated quotation for '{payload.device_name}'")
    return Response(
        content=buf.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="Quotation_{safe_name}.docx"'}
    )
